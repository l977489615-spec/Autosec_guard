#!/usr/bin/env python3
"""Build the <=6000-character conference version using the 41st NCSAC template."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/Users/queen/Desktop/ICV_POC_research/论文/第41次全国计算机安全学术交流会论文模板.docx")
MEDIA = Path("/tmp/paper_v3_media")
OUTPUT = ROOT / "output" / "doc" / "压缩_第41次会议模板修订版.docx"
CORRECTED_FIGURE4 = ROOT / "tmp" / "docs" / "security41_figure4.png"


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_spacing(paragraph, before: float = 0, after: float = 0, line: float = 1.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_paragraph(
    document: Document,
    text: str = "",
    *,
    size: float = 10.5,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_indent: bool = True,
    before: float = 0,
    after: float = 0,
) -> object:
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    set_spacing(paragraph, before, after)
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Pt(size * 2)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)
    return paragraph


def add_labeled_paragraph(document: Document, label: str, text: str, *, size: float = 9) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(paragraph)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size)


def add_heading(document: Document, text: str, level: int) -> None:
    size = 14 if level == 1 else 10.5
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(paragraph, before=4 if level == 1 else 2, after=1)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=True)


def set_cell_text(cell, value: object, *, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(paragraph)
    run = paragraph.add_run(str(value))
    set_run_font(run, 9, bold=bold)


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, attrs in edges.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def add_three_line_table(document: Document, title_cn: str, title_en: str, rows: list[list[object]], widths=None) -> None:
    add_paragraph(
        document,
        title_cn,
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
        before=2,
    )
    add_paragraph(
        document,
        title_en,
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, bold=row_index == 0)
            if widths:
                cell.width = Cm(widths[col_index])
            set_cell_border(
                cell,
                top={"val": "nil"},
                left={"val": "nil"},
                bottom={"val": "nil"},
                right={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )
    for cell in table.rows[0].cells:
        set_cell_border(
            cell,
            top={"val": "single", "sz": "12", "color": "000000"},
            bottom={"val": "single", "sz": "8", "color": "000000"},
        )
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})


def add_figure(document: Document, image: Path, title_cn: str, title_en: str, width_cm: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(paragraph, before=2)
    paragraph.add_run().add_picture(str(image), width=Cm(width_cm))
    add_paragraph(document, title_cn, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_paragraph(document, title_en, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)


def add_equation(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(paragraph, before=1, after=1)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)


def add_reference(document: Document, text: str) -> None:
    paragraph = add_paragraph(document, text, size=9, first_indent=False)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)


def build_corrected_figure4() -> Path:
    from render_baseline_four_metrics_figure import render_figure

    return render_figure(CORRECTED_FIGURE4)


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE}")
    document = Document(TEMPLATE)
    clear_document(document)

    add_paragraph(
        document,
        "证据驱动的车联网多智能体漏洞验证",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )
    add_paragraph(
        document,
        "胡雨翠¹，李奇敖²",
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
        before=3,
    )
    add_paragraph(
        document,
        "（1. 北京天融信网络安全技术有限公司，北京 100193）",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )
    add_paragraph(
        document,
        "（2. 北京航空航天大学网络空间安全学院，北京 100191）",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )

    abstract = (
        "针对智能网联汽车攻击面异构、工具依赖强和证据难复核等问题，提出证据驱动多智能体协同漏洞验证方法EDVV。该方法以证据"
        "缺口表征验证状态，利用检索增强约束PoC选择，依据证据收益、执行成本和风险代价规划路径，在执行失败或证据不足时局部补证。"
        "平台覆盖网络、Android IVI、CAN/UDS/DoIP和无线接口。三类环境实验中，基准阳性召回率、基准子任务完成率和证据完整率"
        "分别为86.7%、93.3%和95.3%，GPT-5.4-mini平均端到端净耗时为5.22 min，体现任务推进与证据复核能力。"
    )
    add_labeled_paragraph(document, "摘  要：", abstract)
    add_labeled_paragraph(document, "关键词：", "智能网联汽车；漏洞验证；多智能体；证据驱动；检索增强")

    add_paragraph(
        document,
        "Evidence-Driven Multi-Agent Vulnerability Verification for Connected Vehicles",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
        before=4,
    )
    add_paragraph(
        document,
        "HU Yucui¹, LI Qiao²",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )
    add_paragraph(
        document,
        "1. Beijing Topsec Network Security Technology Co., Ltd., Beijing 100193, China",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )
    add_paragraph(
        document,
        "2. School of Cyber Science and Technology, Beihang University, Beijing 100191, China",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )
    en_abstract = (
        "To address heterogeneous attack surfaces, dependence on local vehicle-side tools, risky PoC execution, and incomplete evidence, "
        "this paper proposes EDVV, an evidence-driven multi-agent vulnerability verification method. EDVV represents verification states "
        "with evidence gaps, constrains PoC selection through retrieval-augmented knowledge, ranks actions by evidence gain, cost and risk, "
        "and performs local evidence completion after failures. Experiments in three authorized environments achieve 86.7% vulnerability "
        "recall, 93.3% benchmark sub-task completion, and 95.3% evidence completeness. GPT-5.4-mini yields the lowest "
        "mean end-to-end runtime of 5.22 min."
    )
    add_labeled_paragraph(document, "Abstract: ", en_abstract)
    add_labeled_paragraph(document, "Key words: ", "connected vehicle; vulnerability verification; multi-agent; evidence-driven; retrieval augmentation")
    add_paragraph(
        document,
        "作者简介：胡雨翠（1984-），女，湖北洪湖人，工程师，硕士，主要研究方向为信息安全与车联网安全；"
        "李奇敖（2003-），男，云南曲靖人，硕士研究生，主要研究方向为车联网安全。",
        size=9,
        first_indent=False,
        before=3,
    )

    add_heading(document, "1 引言", 1)
    add_paragraph(
        document,
        "智能网联汽车融合车载操作系统、移动应用、无线接口、云端服务和CAN/UDS等车内协议，安全边界扩展至“车-路-云-网-端”。"
        "传统漏洞验证依赖人工选择脚本、配置设备并整理证据，面对网络、Android、蓝牙、Wi-Fi、USB和车内总线等异构攻击面时，"
        "容易出现工具割裂、上下文丢失、高风险动作失控和结果难以复核。PentestGPT、PentestAgent、AutoPenBench等研究证明了"
        "大模型在任务分解、工具调用和长流程推进方面的能力[1-4]，但现有工作主要面向Web、CTF或通用网络服务，对CAN分析仪、"
        "车端接口、安全审批和结构化证据的联合支持不足。"
    )
    add_paragraph(
        document,
        "为此，本文提出证据驱动多智能体协同漏洞验证方法EDVV。其核心不是增加Agent数量，而是以证据缺口驱动状态更新，以证据"
        "收益约束PoC选择，并在授权、能力和风险约束下执行局部补证。主要贡献包括：构建面向车联网的证据驱动验证模型；提出结合"
        "证据完整性、风险和执行成本的路径规划方法；实现侦察、规划、决策、执行、反思和评估协同，以及PoC、ADB和PCAN-USB"
        "等本地工具的受控调用；在仿真、Android IVI和授权实车环境中完成实验评估。"
    )

    add_heading(document, "2 相关工作", 1)
    add_paragraph(
        document,
        "大语言模型在网络安全领域的应用已由安全问答和代码审计扩展到自动化渗透测试。PentestGPT将复杂测试拆分为推理、生成和"
        "解析模块，以缓解长流程中的上下文丢失[1]；PentestAgent引入检索增强与角色协作，组织情报收集、漏洞分析和利用验证[2]。"
        "AutoPenBench、CyBench和HackSynth从任务成功、子任务推进、工具调用和资源消耗等维度评估安全智能体[3-6]。这些工作"
        "表明大模型能够生成计划并根据观察结果调整步骤，但评价目标多为获取flag、完成利用或推进通用渗透任务，对证据归档质量、"
        "高风险动作授权和本地硬件依赖考虑较少。网站漏洞自主利用、多Agent零日利用、CTF智能体及相关综述进一步展示了LLM安全"
        "智能体的能力与评价边界[13-19]。"
    )
    add_paragraph(
        document,
        "车联网安全研究已揭示车载网络、远程接口和无线链路的多类风险。Koscher等验证了现代汽车内部网络的攻击可能性[8]；"
        "Checkoway等系统分析了近程和远程汽车攻击面[9]；Rouf等研究了车内无线网络的隐私与安全问题[10]。后续远程汽车测试平台"
        "和ECU自动模糊测试工作进一步说明，可配置环境和协议工具是提升可复现性的关键[11-12]。然而，现有方法多聚焦单项工具或"
        "单类协议；攻击图、CAN威胁与无线车端攻击研究也表明跨攻击面的统一验证仍是现实需求[20-22]。现有工作缺少将网络侦察、"
        "Android IVI、无线接口、CAN/UDS/DoIP、安全审批和证据留存统一编排的验证方法。"
    )
    add_paragraph(
        document,
        "EDVV借鉴安全智能体的任务分解和反馈思想，同时将验证目标由“完成利用”调整为“检出风险并形成可审计证据”。该差异决定了"
        "系统不仅要判断工具是否执行成功，还要检查响应能否关联目标、证据是否达到预定等级，以及高风险动作是否满足授权边界。"
    )

    add_heading(document, "3 EDVV方法", 1)
    add_heading(document, "3.1 问题定义", 2)
    add_paragraph(
        document,
        "给定测试目标T、PoC库P、检索知识库K、工具集合U和安全策略S，EDVV生成验证计划Plan，并输出执行结果Result、证据集合E"
        "和报告Report。目标上下文C记录资产、协议、接口、设备能力和授权范围；安全策略包含白名单、风险等级、人工审批、沙箱与"
        "审计规则。验证过程建模为有向状态图G=(V,E)，节点表示目标画像、已执行PoC、证据状态和授权状态，边表示侦察、PoC执行、"
        "CAN/UDS测试或补充取证动作。"
    )
    add_paragraph(
        document,
        "EDVV可形式化为六元组<C,K,P,U,S,E>：C为目标上下文，K为漏洞、协议、工具和案例知识，P为PoC插件库，U为本地工具，"
        "S为安全策略，E为证据集合。平台搜索的不是执行数量最多的路径，而是在授权范围内使证据收益较高、风险和人工成本可控的"
        "验证路径。对于每个候选任务，计划项均记录目标、PoC、参数、所需设备、风险等级、预期证据和失败回退策略。"
    )
    add_figure(document, MEDIA / "image1.png", "图1 EDVV总体框架与执行流程", "Fig.1 Overall framework and execution flow of EDVV", 14.5)

    add_heading(document, "3.2 RAG验证决策", 2)
    add_paragraph(
        document,
        "RAG知识库并非简单向模型补充背景，而是把漏洞知识、协议规范、PoC说明、工具文档、历史案例和测试规则转换为可执行约束。"
        "检索输入包括目标类型、开放端口、服务指纹、ADB状态、CAN通道、应用包名、组件版本和已发现线索；输出被整理为候选攻击面、"
        "适用PoC、前置条件、参数要求、风险等级、工具依赖和预期证据。规划Agent据此判断“能否测、测什么、如何测以及应获得何种"
        "证据”，降低脚本名编造、参数错误和不适用PoC调用。"
    )
    add_paragraph(
        document,
        "例如，当侦察结果出现诊断服务或DoIP线索时，知识库补充UDS会话控制、安全访问和响应码解释，并检查PCAN-USB或网络接口"
        "是否可用；当目标为Android IVI时，知识库提供组件导出、WebView、明文通信、敏感存储和第三方库检查规则。检索结果只约束"
        "验证计划，不直接替代工具执行和人工复核。"
    )

    add_heading(document, "3.3 证据评分与约束规划", 2)
    add_paragraph(
        document,
        "对动作e，证据完整性由执行日志Log、截图Screenshot、协议响应Response、文件制品Artifact和审计记录Audit构成。各证据"
        "项满足归档条件时取1，否则取0；权重分别为0.15、0.15、0.30、0.25和0.15。协议响应与文件制品可直接支撑漏洞真实性复核，"
        "因此权重较高。"
    )
    add_equation(document, "EvidenceScore(e) = Σ wᵢeᵢ，Σwᵢ = 1                                      （1）")
    add_three_line_table(
        document,
        "表1 证据链等级",
        "Tab.1 Evidence chain levels",
        [
            ["等级", "证据组成", "用途"],
            ["L1", "日志", "证明动作已执行"],
            ["L2", "日志+截图", "复核界面或扫描现象"],
            ["L3", "L2+协议响应", "确认目标交互结果"],
            ["L4", "L3+文件制品", "支持复现和离线检查"],
            ["L5", "L4+审计记录", "报告、复核与合规审计"],
        ],
        widths=[2.0, 6.0, 6.0],
    )
    add_paragraph(
        document,
        "证据等级用于区分“存在记录”和“足以支撑结论”。若候选漏洞仅获得L1或L2证据，而任务要求协议响应或文件制品，评估Agent"
        "不会直接输出完整验证结论，而是把缺失项写入状态图并触发补证。各项权重在实验前固定，不根据结果调整，避免事后调参放大"
        "收益。L3加入协议响应，可确认目标确实产生可解释交互；L4进一步保存配置、版本或样本制品，支持离线检查与复现；L5在L4"
        "基础上加入授权审批和工具调用审计记录，形成可用于正式报告、人工复核和合规审计的完整证据闭环。"
    )
    add_paragraph(
        document,
        "动作收益综合漏洞验证可能性、证据评分和新增攻击面覆盖价值；路径效用在累计证据收益基础上扣除执行成本与风险代价。安全"
        "控制器先移除不满足白名单、授权或设备能力的动作，决策Agent再从剩余候选中选择效用较高的路径。"
    )
    add_equation(document, "U(π) = Σq(eᵢ) - λΣc(eᵢ) - μΣr(eᵢ) + ηB(π)                       （2）")
    add_heading(document, "3.4 多智能体协同与局部补证", 2)
    add_paragraph(
        document,
        "侦察Agent生成资产、端口、服务、ADB状态和车内协议线索；规划Agent结合检索知识生成候选任务；决策Agent依据证据收益、"
        "风险和能力约束选择路径；执行Agent仅调用通过安全门控的工具；评估Agent判定风险与证据等级；反思Agent根据参数缺失、"
        "工具失败或证据不足触发局部重规划。Agent之间传递结构化状态而非自由文本命令，使每次工具调用均可关联任务标识、参数、"
        "返回码、证据文件和授权记录。"
    )
    add_paragraph(
        document,
        "完整流程包括六个步骤：①构建目标上下文并确认授权范围；②检索漏洞、协议和工具知识；③生成候选任务并计算证据收益、成本"
        "和风险；④安全控制器剪除不满足白名单、授权和能力约束的动作；⑤执行PoC或本地工具并归档结果；⑥若证据不足或执行失败，"
        "仅对受影响路径补侦察、补参数、切换工具或补充取证，否则生成报告。"
    )
    add_figure(document, MEDIA / "image2.png", "图2 证据驱动规划与局部补证流程", "Fig.2 Evidence-driven planning and local evidence completion", 14.5)
    add_figure(document, MEDIA / "image3.png", "图3 EDVV评分驱动的多智能体协同机制", "Fig.3 Score-driven multi-agent collaboration in EDVV", 14.5)

    add_heading(document, "4 平台实现", 1)
    add_paragraph(
        document,
        "平台采用“控制台-任务引擎-多智能体编排-工具适配-证据归档”结构。PoC插件通过统一元数据描述攻击面、目标系统、输入参数、"
        "风险等级、硬件依赖、授权要求和证据提取方式；检索知识库保存漏洞知识、协议规范、工具说明和历史案例；执行器将结构化计划"
        "映射为网络扫描、ADB、蓝牙/Wi-Fi、USB、CAN/UDS/DoIP等调用。高风险PoC执行前检查目标白名单、授权状态和风险等级，"
        "沙箱限制CPU、内存、输出和访问目标，所有结果保存为结构化记录。"
    )
    add_paragraph(
        document,
        "当前工具库包含146项可执行PoC：侦察8项、网络服务15项、应用安全46项、系统配置12项、第三方组件5项、车内协议13项、"
        "无线与外设37项、高级攻击10项。CAN联动通过PCAN-USB记录CAN ID、帧类型、发送次数、UDS响应、异常现象和证据文件，使"
        "外部攻击面发现能够延伸至车内协议验证。"
    )
    add_three_line_table(
        document,
        "表2 PoC攻击面构成",
        "Tab.2 Composition of the PoC attack surfaces",
        [
            ["攻击面", "数量", "典型检测项", "硬件依赖"],
            ["侦察", "8", "存活、端口、服务识别", "否"],
            ["网络服务", "15", "ADB、SSH、FTP、RTSP", "部分"],
            ["应用与系统", "58", "组件、WebView、系统配置", "ADB/样本"],
            ["第三方组件", "5", "OpenSSL、libpng等", "文件"],
            ["车内协议", "13", "CAN、UDS、DoIP", "PCAN-USB"],
            ["无线与外设", "37", "蓝牙、Wi-Fi、USB、SDR", "是"],
            ["高级攻击", "10", "OTA、RF、GPS、固件", "部分"],
            ["合计", "146", "-", "-"],
        ],
        widths=[3.1, 1.5, 6.8, 2.8],
    )
    add_paragraph(
        document,
        "PoC元数据驱动执行链先完成侦察，再根据目标接口、开放服务、required_params和profile筛选候选脚本。执行器对参数进行类型"
        "与范围校验，在独立子进程中限制CPU、内存、文件句柄、输出大小和可访问目标；需要人工审批的任务在终端请求确认，等待时间"
        "单独记录，不计入自动化净耗时。执行完成后统一写入状态、漏洞判断、错误、日志、原始响应和证据路径。"
    )

    add_heading(document, "5 实验设计", 1)
    add_heading(document, "5.1 环境与评价指标", 2)
    add_paragraph(
        document,
        "实验覆盖受控仿真环境、Android/Linux IVI设备和大众ID.4授权实车。仿真环境验证网络服务、审批和证据归档；Android IVI"
        "通过WLAN/IP与ADB/USB验证组件暴露、WebView、敏感数据和系统配置；实车通过WLAN/IP、蓝牙、Wi-Fi和PCAN-USB验证网络、"
        "无线及CAN/UDS/DoIP任务。三类环境使用相同PoC元数据、安全规则和证据字段。"
    )
    add_paragraph(
        document,
        "基准阳性集合由受控正样本、基准扫描和人工复核共同确定，并按唯一PoC去重。全文采用统一的四项核心指标，计算口径如下。"
    )
    add_paragraph(
        document,
        "基准阳性召回率（Vulnerability Recall） = Agent检出的基准阳性PoC数 / 基准阳性PoC总数 × 100%。分母是经受控正样本、基准扫描或人工复核确认确实"
        "存在风险的唯一PoC数，表示实验中应当被平台发现的客观阳性集合；分子是Agent实际选择并完成执行，且风险结论与人工复核"
        "一致的基准阳性PoC数。仅生成计划、启动后报错、未获得目标响应或结论无法确认的PoC均不计入分子。该指标越高，说明平台"
        "找回真实风险的能力越强，漏检越少。"
    )
    add_paragraph(
        document,
        "基准子任务完成率（Benchmark Sub-task Completion Rate） = 已完成基准子任务数 / 基准子任务总数 × 100%。分母不是PoC库中的全部146项脚本，而是依据目标接口、"
        "设备能力、开放服务和执行前置条件确定的30项可执行子任务，表示本轮实验应当推进的标准任务集合；分子是已经完成至少一次"
        "授权执行、获得可解析结果并建立证据索引的子任务数。目标不具备相关接口或硬件能力的任务不进入分母，人工明确拒绝的高风险"
        "任务也不作为自动化失败。该指标越高，说明智能体将计划推进到执行和结果归档阶段的能力越强。"
    )
    add_paragraph(
        document,
        "证据完整率（Evidence Completeness Rate） = 达到L5证据等级的已执行PoC数 / 已执行PoC总数 × 100%。分母是进入执行器并形成明确终止状态与结构化结果的"
        "唯一PoC数；分子是其中证据链达到表1所定义L5等级的PoC数。达到L5意味着该执行项同时具备日志、截图、协议响应、文件制品"
        "和审计记录，能够完整说明执行了什么、目标如何响应、结论依据是什么以及执行是否经过授权。仅有日志、错误信息或结构化返回，"
        "但缺少上述任一证据层级的执行均不计入分子。后文各表及图中的证据完整率均采用该L5口径。"
    )
    add_paragraph(
        document,
        "平均端到端净耗时（Mean End-to-End Runtime） = Σ（任务结束时间 - 任务开始时间 - 人工等待时间）/ 有效实验轮次。计时从智能体接收目标并开始侦察时"
        "起，到最后一个验证步骤结束且结构化结果完成归档时止，包含规划、模型调用、PoC执行、反思补证和报告整理时间；等待操作者"
        "输入授权结果的时间单独记录并扣除。该指标用于衡量完整自动化验证闭环的实际效率，数值越低表示处理速度越快。"
    )
    add_paragraph(
        document,
        "同一PoC的重试、反思重入和跨目标重复执行在综合统计中只计一次，并采用最终有效结果。该去重规则避免重复运行扩大已执行数，"
        "同时避免把目标不具备的ADB、无线或车内协议能力错误计为漏检。漏报率作为基准阳性召回率的补充指标，按未检出的基准阳性PoC数"
        "除以基准阳性PoC总数计算。"
    )
    add_heading(document, "5.2 验证闭环结果", 2)
    add_three_line_table(
        document,
        "表3 EDVV验证闭环结果",
        "Tab.3 Closed-loop verification results of EDVV",
        [
            ["指标", "原始计数", "结果"],
            ["基准阳性召回率", "26/30", "86.7%"],
            ["基准子任务完成率", "28/30", "93.3%"],
            ["证据完整率（L5）", "61/64", "95.3%"],
            ["漏报率", "4/30", "13.3%"],
        ],
        widths=[6.0, 4.0, 4.0],
    )
    add_paragraph(
        document,
        "146项PoC中有124项在相应目标条件下完成授权执行。30项去重基准阳性中，EDVV正确检出26项并完成28项基准验证子任务。"
        "漏报主要来自车内协议前置状态、网络服务响应不稳定和应用侧证据不足。64项已执行PoC中有61项形成L5证据链；其余3项虽形成"
        "明确执行结果，但未同时具备日志、截图、协议响应、文件制品和审计记录，因此未计入证据完整率分子。"
    )
    add_paragraph(
        document,
        "分类结果显示，侦察、无线、系统配置和高级攻击类基准阳性召回较稳定；车内协议和应用安全类仍存在漏报。前者主要受诊断会话、"
        "总线状态和DoIP响应条件影响，后者与应用样本、ADB权限及静态制品可用性有关。这说明PoC库规模不能直接代表验证能力，目标"
        "适配、前置状态和证据质量同样决定最终结果。"
    )

    add_heading(document, "6 结果分析", 1)
    add_heading(document, "6.1 协同机制消融", 2)
    add_three_line_table(
        document,
        "表4 协同机制消融结果",
        "Tab.4 Ablation results of collaborative mechanisms",
        [
            ["配置", "基准阳性召回率/%", "基准子任务完成率/%", "证据完整率/%", "平均端到端净耗时/min"],
            ["单智能体", "23.3", "23.3", "100.0", "1.95"],
            ["多智能体", "66.7", "70.0", "62.5", "10.26"],
            ["多智能体+反思", "80.0", "83.3", "89.1", "13.46"],
            ["EDVV", "86.7", "93.3", "95.3", "14.73"],
        ],
        widths=[5.2, 2.4, 2.4, 2.8, 2.4],
    )
    add_paragraph(
        document,
        "单智能体执行项较少，因此其证据完整率较高不能代表总体验证能力；完整EDVV显著提高基准阳性召回率和基准子任务完成率，表明"
        "角色分工、检索约束和反思补证有助于长流程推进。"
    )
    add_heading(document, "6.2 基线与模型对比", 2)
    add_three_line_table(
        document,
        "表5 证据评分候选选择策略对比",
        "Tab.5 Comparison of candidate selection strategies",
        [
            ["策略", "排序依据", "基准子任务完成率/%", "证据完整率/%"],
            ["随机选择PoC", "从候选PoC中随机选择", "70.0", "62.0"],
            ["成功率优先", "优先选择历史执行成功率较高的PoC", "80.0", "71.0"],
            ["EDVV证据收益排序", "综合证据评分、覆盖价值、风险与成本", "90.0", "79.0"],
        ],
        widths=[3.5, 6.3, 2.8, 2.4],
    )
    add_paragraph(
        document,
        "候选策略对比表明，仅依据历史成功率能够提高可执行性，但不能保证形成高等级证据；EDVV将证据缺口、攻击面覆盖、风险代价"
        "和执行成本共同纳入排序，因此在基准子任务完成率和证据完整率上均优于随机选择与成功率优先策略。"
    )
    add_figure(
        document,
        build_corrected_figure4(),
        "图4 基线方法四指标对比",
        "Fig.4 Comparison of four metrics among baseline methods",
        14.5,
    )
    add_paragraph(
        document,
        "PentestGPT基线与EDVV共享受控执行器和评价口径，差异主要在规划、证据要求和补证机制。EDVV相较PentestGPT将基准阳性召回率"
        "提高40.0个百分点、基准子任务完成率提高33.3个百分点、证据完整率提高36.6个百分点。普通多智能体用于展示角色分工带来的"
        "阶段性变化。"
    )
    add_three_line_table(
        document,
        "表6 不同基础模型的EDVV结果",
        "Tab.6 EDVV results with different foundation models",
        [
            ["模型", "基准阳性召回率/%", "基准子任务完成率/%", "证据完整率/%", "平均端到端净耗时/min"],
            ["DeepSeek v4 pro", "86.7", "90.0", "98.7", "32.52"],
            ["智谱GLM-5", "86.7", "93.3", "95.3", "14.73"],
            ["GPT-5.4-mini", "83.3", "93.3", "97.6", "5.22"],
            ["千问qwen-max", "86.7", "93.3", "90.0", "8.18"],
        ],
        widths=[5.2, 2.4, 2.4, 2.8, 2.4],
    )
    add_paragraph(
        document,
        "四种模型均可驱动相同PoC和工具接口。GLM-5、DeepSeek和qwen-max的基准阳性召回率均为86.7%；GPT-5.4-mini为83.3%，"
        "但平均端到端净耗时最低。结果说明平台能力来自证据驱动流程与工具约束，而非依赖单一模型。"
    )
    add_paragraph(
        document,
        "模型对比采用相同PoC库、任务集合、工具接口和安全策略。DeepSeek形成的验证计划更长、耗时较高；GPT-5.4-mini具有明显"
        "时延优势；GLM-5在基准阳性召回率、基准子任务完成率和平均端到端净耗时之间取得较均衡结果。"
    )
    add_heading(document, "6.3 安全控制与车端联动", 2)
    add_paragraph(
        document,
        "146项PoC按Low、Medium、High和Critical分级，其中10项要求授权。实验中1项未授权动作被阻断，123项在授权条件下执行，"
        "审计记录完整率为100%。PCAN-USB联动完成12项CAN/UDS相关验证，覆盖CAN重放、UDS SecurityAccess、DoIP探测和受控模糊"
        "测试。高风险动作执行前均检查白名单、设备能力和授权状态，人工等待时间不计入端到端净耗时。"
    )
    add_paragraph(
        document,
        "安全控制采用“元数据预判-策略检查-人工授权-沙箱执行-审计归档”五层流程。Low和Medium任务在参数完备时可自动执行；"
        "High和Critical任务根据破坏级别、目标白名单和授权策略决定是否请求确认。被拒绝或能力不满足的任务记录为跳过或阻断，不"
        "计入自动化任务失败。该口径区分了模型规划失败、工具执行失败和操作者主动拒绝，避免高风险人工跳过拉低系统能力指标。"
    )
    add_heading(document, "6.4 有效性威胁", 2)
    add_paragraph(
        document,
        "实验仅覆盖三类目标和有限车型版本，基准阳性集合仍可能遗漏依赖特定硬件状态或长期触发条件的风险；不同模型的单轮结果存在"
        "采样波动；中间配置的实验规模仍需扩大。后续需增加车型、独立重复运行和置信区间，并由双人复核证据等级，以进一步验证"
        "泛化能力。"
    )

    add_heading(document, "7 结论", 1)
    add_paragraph(
        document,
        "本文提出EDVV方法，以证据缺口、证据评分和风险约束组织车联网漏洞验证，通过多智能体协同与局部补证将PoC执行转化为可"
        "复核的验证闭环。平台支持146项PoC及网络、Android IVI、无线和CAN/UDS/DoIP工具联动。实验中EDVV取得86.7%的基准"
        "阳性召回率、93.3%的基准子任务完成率和95.3%的证据完整率，并能对高风险动作实施授权与审计。后续将补充独立消融实验、"
        "多车型验证和更细粒度的证据自动评定。"
    )

    add_heading(document, "参考文献", 1)
    references = [
        "[1] DENG G, LIU Y, MAYORAL-VILCHES V, et al. PentestGPT: Evaluating and harnessing large language models for automated penetration testing[C]//USENIX Security. 2024: 847-864.",
        "[2] SHEN X, WANG L, LI Z, et al. PentestAgent: Incorporating LLM agents to automated penetration testing[EB/OL]. arXiv:2411.05185, 2024.",
        "[3] GIOACCHINI L, MELLIA M, DRAGO I, et al. AutoPenBench: Benchmarking generative agents for penetration testing[EB/OL]. arXiv:2410.03225, 2024.",
        "[4] ZHANG A K, PERRY N, DULEPET R, et al. Cybench: A framework for evaluating cybersecurity capabilities and risks of language models[EB/OL]. arXiv:2408.08926, 2024.",
        "[5] BHATT M, CHENNABASAPPA S, NIKOLAIDIS C, et al. CyberSecEval 2: A wide-ranging cybersecurity evaluation suite for large language models[EB/OL]. arXiv:2404.13161, 2024.",
        "[6] MUZSAI L, IMOLAI D, LUKACS A. HackSynth: LLM agent and evaluation framework for autonomous penetration testing[EB/OL]. arXiv:2412.01778, 2024.",
        "[7] YAO S, ZHAO J, YU D, et al. ReAct: Synergizing reasoning and acting in language models[C]//ICLR. 2023.",
        "[8] KOSCHER K, CZESKIS A, ROESNER F, et al. Experimental security analysis of a modern automobile[C]//IEEE Symposium on Security and Privacy. 2010: 447-462.",
        "[9] CHECKOWAY S, MCCOY D, KANTOR B, et al. Comprehensive experimental analyses of automotive attack surfaces[C]//USENIX Security. 2011.",
        "[10] ROUF I, MILLER R, MUSTAFA H, et al. Security and privacy vulnerabilities of in-car wireless networks[C]//USENIX Security. 2010.",
        "[11] KULANDAIVEL S, LU W, BARRY B, et al. Towards a configurable and practical remote automotive security testing platform[EB/OL]. arXiv:2404.02291, 2024.",
        "[12] WERQUIN T, HUBRECHTSEN R, THANGARAJAN A, et al. Automated fuzzing of automotive control units[EB/OL]. arXiv:2102.12345, 2021.",
        "[13] FANG R, BINDU R, GUPTA A, et al. LLM agents can autonomously hack websites[EB/OL]. arXiv:2402.06664, 2024.",
        "[14] FANG R, BINDU R, GUPTA A, et al. LLM agents can autonomously exploit one-day vulnerabilities[EB/OL]. arXiv:2404.08144, 2024.",
        "[15] FANG R, BINDU R, GUPTA A, et al. Teams of LLM agents can exploit zero-day vulnerabilities[EB/OL]. arXiv:2406.01637, 2024.",
        "[16] ABRAMOVICH G, DENG G, GANESH S, et al. EnIGMA: Enhanced interactive generative model agent for CTF challenges[EB/OL]. arXiv:2409.16165, 2024.",
        "[17] A survey of large language models in cybersecurity[EB/OL]. arXiv:2402.16968, 2024.",
        "[18] Large language models for cyber security: A systematic literature review[EB/OL]. arXiv:2405.04760, 2024.",
        "[19] CritBench: A framework for evaluating cybersecurity capabilities of large language models in IEC 61850 digital substation environments[EB/OL]. arXiv:2604.06019, 2026.",
        "[20] SAULAIMAN M N E, CSILLING A, KOZLOVSZKY M. Integrated automation for threat analysis and risk assessment in automotive cybersecurity through attack graphs[J]. Acta Polytechnica Hungarica, 2025, 22(2): 149-168.",
        "[21] HOPPE T, KILTZ S, DITTMANN J. Security threats to automotive CAN networks: Practical examples and selected short-term countermeasures[J]. Reliability Engineering & System Safety, 2011, 96(1): 11-25.",
        "[22] WOO S, JO H J, LEE D H. A practical wireless attack on the connected car and security protocol for in-vehicle CAN[J]. IEEE Transactions on Intelligent Transportation Systems, 2015, 16(2): 546-556.",
    ]
    for reference in references:
        add_reference(document, reference)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)

    check = Document(OUTPUT)
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", "\n".join(p.text for p in check.paragraphs)))
    total_chars_without_refs = 0
    for paragraph in check.paragraphs:
        if paragraph.text.strip() == "参考文献":
            break
        total_chars_without_refs += len(re.sub(r"\s+", "", paragraph.text))
    if chinese_chars > 6000:
        raise RuntimeError(f"Chinese character count exceeds 6000: {chinese_chars}")
    print(f"{OUTPUT}\nChinese characters: {chinese_chars}\nMain text chars before references: {total_chars_without_refs}\nTables: {len(check.tables)}")


if __name__ == "__main__":
    main()
