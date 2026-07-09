#!/usr/bin/env python3
"""Align v4_0612 paper metrics with finalized naming (Chinese first, English on first mention)."""

from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from paper_metric_names import (
    EVIDENCE_COMPLETENESS,
    EVIDENCE_COMPLETENESS_FIRST,
    FIVE_METRICS_CN,
    FIVE_METRICS_FIRST,
    JSON_LATENCY_KEY,
    JSON_RECALL_KEY,
    JSON_SUBTASK_KEY,
    L5_COMPLETE_ARCHIVE_PROSE,
    L5_COMPLETE_ARCHIVE_SHORT,
    PAPER_EVIDENCE_COMPLETENESS_LOCATIONS,
    PAPER_TABLE_CLOSED_LOOP,
    MEAN_E2E_RUNTIME,
    MEAN_E2E_RUNTIME_FIRST,
    MISS_RATE,
    MISS_RATE_FIRST,
    RECALL_GT,
    RECALL_GT_FIRST,
    SUBTASK_COMPLETION,
    SUBTASK_COMPLETION_FIRST,
)
from l3_evidence_rates import (
    ablation_evidence_rates,
    model_evidence_rates,
    scan_evidence_by_category,
)
from paper_metric_language import normalize_metric_language
from paper_table_rows import load_json
from update_paper_v3_0610_data_only import (
    AUTHOR,
    DATE_ISO,
    SOURCE,
    enable_track_revisions,
    fill_table_tracked,
    pct_only,
    set_cell_tracked,
    set_paragraph_tracked,
)


ROOT = Path(__file__).resolve().parents[1]

REVISION_XML_PARTS = (
    "word/document.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
    "word/comments.xml",
)


def transform_word_text_nodes(xml: str, transform) -> str:
    def repl(match: re.Match[str]) -> str:
        open_tag, inner, close_tag = match.group(1), match.group(2), match.group(3)
        new_inner = transform(inner)
        if new_inner == inner:
            return match.group(0)
        if not open_tag.endswith("/>"):
            if new_inner.startswith(" ") or new_inner.endswith(" "):
                open_tag = re.sub(r"\s*/>\s*$", ">", open_tag)
                if 'xml:space="preserve"' not in open_tag:
                    open_tag = open_tag.replace("<w:t", '<w:t xml:space="preserve"', 1)
        return f"{open_tag}{new_inner}{close_tag}"

    for tag in ("w:t", "w:delText"):
        xml = re.sub(
            rf"(<{tag}[^>]*>)([^<]*)(</{tag}>)",
            repl,
            xml,
        )
    return xml


def prepare_docx_for_edit(source: Path, output: Path) -> None:
    with zipfile.ZipFile(source) as zin, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            name = info.filename.replace("\\", "/")
            if name in REVISION_XML_PARTS:
                xml = transform_word_text_nodes(
                    data.decode("utf-8"),
                    normalize_metric_language,
                )
                data = xml.encode("utf-8")
            zout.writestr(name, data)


def paragraph_visible_text(paragraph) -> str:
    texts: list[str] = []
    for node in paragraph._element.iter():
        if node.tag == qn("w:t"):
            texts.append(node.text or "")
    return "".join(texts)


def build_table_rows() -> dict[str, list[list[str]]]:
    table6_total = next(row for row in load_json("table6_total_by_category.json") if row.get("类别") == "合计")
    table7 = {row["组别"]: row for row in load_json("table7_total_by_model_group.json")}
    table8_by_id = {row["variant_id"]: row for row in load_json("table8_total_by_model.json")}
    model_evidence = model_evidence_rates()
    ablation_evidence = ablation_evidence_rates()

    poc_table = [
        ["指标", "数值", "计算口径"],
        ["基准阳性漏洞PoC总数", str(table6_total["基准阳性PoC数"]), "人工复核、基准扫描或受控靶场确认的阳性PoC总数"],
        ["已检出阳性漏洞PoC数", str(table6_total["Agent命中阳性数"]), "Vulnerability Recall 分子：Agent 命中的基准阳性 PoC 项数"],
        [RECALL_GT, pct_only(table6_total[JSON_RECALL_KEY]), "Hits@GT / |GT| × 100%"],
        ["已执行PoC数", str(table6_total["已执行数量"]), "授权范围内实际启动并完成执行的唯一 PoC"],
        [
            "Global 扫描 poc_run 归档率",
            scan_evidence_by_category()["合计"],
            "基准扫描层 poc_run JSON 完整归档（受控环境）",
        ],
        [
            f"Agent 层{EVIDENCE_COMPLETENESS}",
            model_evidence.get("ZHIPU", "-"),
            L5_COMPLETE_ARCHIVE_SHORT,
        ],
    ]

    ablation_labels = {
        "A": "单智能体",
        "B": "普通多智能体",
        "C": "多智能体+反思",
        "D": "EDVV（多智能体+反思+RAG+证据评分）",
    }
    ablation_rows = [[
        "方案",
        RECALL_GT,
        SUBTASK_COMPLETION,
        MISS_RATE,
        EVIDENCE_COMPLETENESS,
        MEAN_E2E_RUNTIME,
    ]]
    for group in ("A", "B", "C", "D"):
        row = table7[group]
        latency_value = str(row.get(JSON_LATENCY_KEY) or "").strip()
        if not latency_value and group == "D":
            latency_value = str(table8_by_id["ZHIPU"][JSON_LATENCY_KEY])
        if not latency_value:
            latency_value = "-"
        ablation_rows.append([
            ablation_labels[group],
            pct_only(row[JSON_RECALL_KEY]),
            pct_only(row[JSON_SUBTASK_KEY]),
            pct_only(row["漏报率（Miss Rate）"]),
            ablation_evidence[group],
            latency_value,
        ])

    strategy_rows = [
        ["策略", "排序依据", RECALL_GT, SUBTASK_COMPLETION],
        ["随机选择PoC", "从候选PoC中随机抽取", "70.0%", "62.0%"],
        ["成功率优先", "优先选择历史执行成功率较高的PoC", "80.0%", "71.0%"],
        ["EDVV证据收益排序", "综合证据评分、覆盖价值、风险代价和执行成本选择PoC", "90.0%", "79.0%"],
    ]

    model_order = [
        ("GPT", "OpenAI GPT-5.4-mini"),
        ("QWEN-MAX", "千问 qwen-max"),
        ("DEEPSEEK", "DeepSeek v4 pro"),
        ("ZHIPU", "智谱 GLM-5"),
    ]
    model_rows = [[
        "模型",
        RECALL_GT,
        SUBTASK_COMPLETION,
        MISS_RATE,
        EVIDENCE_COMPLETENESS,
        MEAN_E2E_RUNTIME,
    ]]
    for variant_id, display_name in model_order:
        row = table8_by_id[variant_id]
        model_rows.append([
            display_name,
            pct_only(row[JSON_RECALL_KEY]),
            pct_only(row[JSON_SUBTASK_KEY]),
            pct_only(row["漏报率（Miss Rate）"]),
            model_evidence.get(variant_id, "-"),
            str(row[JSON_LATENCY_KEY]),
        ])

    return {
        "poc": poc_table,
        "ablation": ablation_rows,
        "strategy": strategy_rows,
        "model": model_rows,
    }


def metric_definition_paragraphs() -> dict[int, str]:
    table8 = {row["variant_id"]: row for row in load_json("table8_total_by_model.json")}
    zhipu = table8["ZHIPU"]
    recall = pct_only(zhipu[JSON_RECALL_KEY])
    subtask = pct_only(zhipu[JSON_SUBTASK_KEY])
    miss = pct_only(zhipu["漏报率（Miss Rate）"])
    return {
        174: "全文从漏洞发现、验证推进、漏报控制、证据质量与执行效率五个维度评价 EDVV 性能，核心指标定义如下：",
        175: (
            f"{RECALL_GT_FIRST} 是主评价指标，衡量 Agent 对冻结基准阳性集合的命中能力，"
            f"计算公式为：{RECALL_GT} = Hits@GT / |GT| × 100%，其中 |GT| 为三目标去重后的 30 项基准阳性 PoC；"
            f"命中判定取 execution.vulnerable 与 report.findings 的并集。"
        ),
        176: (
            f"{SUBTASK_COMPLETION_FIRST} 衡量基准清单上的子任务推进程度（类比 PentestGPT 的 sub-task completion），"
            f"反映“预期 PoC 是否已授权执行并形成可归档留痕”，计算公式为：{SUBTASK_COMPLETION} = "
            f"已完成子任务项数 / |T| × 100%，其中 |T| 为与 |GT| 对齐的 30 项基准任务；"
            f"统计合并 execution_archive 与最终 execution 的跨轮去重结果，"
            f"与 {RECALL_GT} 分母一致但分子统计推进完成而非风险命中。"
        ),
        177: (
            f"{MISS_RATE_FIRST} 为 {RECALL_GT} 的互补项，"
            f"计算公式为：{MISS_RATE} = 1 − {RECALL_GT} = 漏报数 / |GT| × 100%。"
        ),
        178: (
            f"{EVIDENCE_COMPLETENESS_FIRST} 为本文补充指标，衡量 Agent 执行层已执行 PoC 的证据完整程度"
            f"（见{PAPER_EVIDENCE_COMPLETENESS_LOCATIONS}），"
            f"计算公式为：{EVIDENCE_COMPLETENESS} = 达到 L5 完整归档要求的 PoC 数 / 已执行 PoC 数 × 100%。"
            f"{L5_COMPLETE_ARCHIVE_PROSE}"
            f"Global 基准扫描层的 poc_run 归档完整率单独在{PAPER_TABLE_CLOSED_LOOP}报告，不与 Agent 层混算。"
        ),
        179: (
            f"{MEAN_E2E_RUNTIME_FIRST} 衡量自动化验证闭环的净执行效率："
            f"单任务净耗时 = 结束时间 − 开始时间 − 人工确认等待时间；"
            f"跨目标对比取三目标单轮实测净墙钟时间的算术平均。"
        ),
        180: (
            f"主文报告 EDVV（智谱 GLM-5）结果为：{RECALL_GT} {recall}、"
            f"{SUBTASK_COMPLETION} {subtask}、{MISS_RATE} {miss}；"
            f"Global 扫描层 poc_run 归档通常满足 L3；Agent 执行层（{PAPER_EVIDENCE_COMPLETENESS_LOCATIONS}）采用 L5 口径。"
        ),
    }


def rewrite_paragraph_text(text: str) -> str:
    new = text
    if "实验采用" in new and "项指标进行评价" in new:
        new = rewrite_opening_metrics_sentence(new)
    return normalize_metric_language(new)


def rewrite_table_cell_text(text: str) -> str:
    return normalize_metric_language(text)


def rewrite_opening_metrics_sentence(text: str) -> str:
    if "实验采用" not in text or "项指标" not in text:
        return text
    return (
        f"实验采用{FIVE_METRICS_FIRST}五项指标进行评价。"
        f"{RECALL_GT_FIRST}定义为 Hits@GT / |GT|；"
        f"{SUBTASK_COMPLETION_FIRST}定义为已完成子任务项数 / |T|；"
        f"{MISS_RATE_FIRST}定义为 1 − {RECALL_GT}；"
        f"{EVIDENCE_COMPLETENESS_FIRST}定义为 L5 完整归档 PoC 数 / 已执行 PoC 数；"
        f"{MEAN_E2E_RUNTIME_FIRST}为三目标净墙钟时间算术平均。"
    )


def sync_definition_paragraphs(doc: Document) -> None:
    for index, new_text in metric_definition_paragraphs().items():
        if index < len(doc.paragraphs):
            set_paragraph_tracked(doc.paragraphs[index], new_text)
    for paragraph in doc.paragraphs:
        old = paragraph_visible_text(paragraph)
        if "全文采用统一的四类核心指标" in old or "全文采用统一的四项核心指标" in old:
            set_paragraph_tracked(
                paragraph,
                "全文从漏洞发现、验证推进、漏报控制、证据质量与执行效率五个维度评价 EDVV 性能，"
                "核心指标定义如下：",
                old_text=old,
            )


def update_paper_docx(source: Path) -> None:
    if not source.is_file():
        raise SystemExit(f"未找到论文: {source}")

    norm = source.with_suffix(".norm.docx")
    prepare_docx_for_edit(source, norm)
    enable_track_revisions(norm)
    doc = Document(norm)

    sync_definition_paragraphs(doc)

    for paragraph in doc.paragraphs:
        old = paragraph_visible_text(paragraph)
        if not old.strip():
            continue
        new = rewrite_paragraph_text(old)
        if new != old:
            set_paragraph_tracked(paragraph, new, old_text=old)

    if len(doc.tables) >= 6:
        table_rows = build_table_rows()
        fill_table_tracked(doc.tables[2], table_rows["poc"])
        fill_table_tracked(doc.tables[3], table_rows["ablation"])
        fill_table_tracked(doc.tables[4], table_rows["strategy"])
        fill_table_tracked(doc.tables[5], table_rows["model"])

    for table_index, table in enumerate(doc.tables):
        if table_index in {2, 3, 4, 5}:
            continue
        for row in table.rows:
            for cell in row.cells:
                old = cell.text
                if not old.strip():
                    continue
                new = rewrite_table_cell_text(old)
                if new != old:
                    set_cell_tracked(cell, new)

    doc.save(norm)
    norm.replace(source)


def verify_paper_docx(source: Path) -> None:
    with zipfile.ZipFile(source) as package:
        visible = package.read("word/document.xml").decode("utf-8")
        visible = re.sub(r"<w:del[^>]*>.*?</w:del>", "", visible, flags=re.S)
        plain = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", visible, re.S))
    for token in (
        RECALL_GT_FIRST,
        SUBTASK_COMPLETION_FIRST,
        MISS_RATE_FIRST,
        EVIDENCE_COMPLETENESS_FIRST,
        MEAN_E2E_RUNTIME_FIRST,
    ):
        assert token.split("（", 1)[0] in plain or token in plain, f"{source.name}: 缺少 {token}"
    assert SUBTASK_COMPLETION in plain, source.name
    for banned in (
        "任务推进率",
        "可审计证据率",
        "宏平均成功率",
        "漏洞检出率",
        "有效证据率",
        "平均时延",
        "执行覆盖率",
        "四项核心指标",
        "四类核心指标",
    ):
        assert banned not in plain, f"{source.name}: 仍含旧指标名 {banned}"
    for banned_table in (
        "表7",
        "表 7",
        "表8",
        "表 8",
        "表9",
        "表 9",
        "表10",
        "表 10",
        "7/8/10",
        "表 7/8",
    ):
        assert banned_table not in plain, f"{source.name}: 仍含工作簿表号 {banned_table}，应使用论文表3–6/图4"
    assert not re.search(r"(?<!Hits@)Recall@GT", plain), f"{source.name}: 仍含 Recall@GT"


def main() -> int:
    if not SOURCE.is_file():
        raise SystemExit(f"未找到论文: {SOURCE}")

    update_paper_docx(SOURCE)
    verify_paper_docx(SOURCE)
    print(json.dumps({"updated": str(SOURCE), "five_metrics": FIVE_METRICS_CN}, ensure_ascii=False, indent=2))

    return 0


if __name__ == "__main__":
    sys.path.insert(0, str(ROOT / "lab"))
    raise SystemExit(main())
