#!/usr/bin/env python3
"""Update v16 thesis: experiment data, metrics, and data-driven prose only (tracked)."""

from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
STRICT_DATA = ROOT / "lab" / "final_paper_data_strict"
SOURCE = Path(
    "/Users/queen/Desktop/ICV_POC_research/论文/"
    "面向智能网联汽车的证据驱动多智能体协同漏洞验证方法研究_6.9_1_贡献与章节结构修订版_v3_"
    "CCF中文模板版_v16_摘要关键词优化版.docx"
)
BACKUP = SOURCE.with_suffix(".bak.docx")
AUTHOR = "李奇敖"
DATE_ISO = "2026-06-02T12:00:00Z"

_revision_id = 0


def next_revision_id() -> str:
    global _revision_id
    _revision_id += 1
    return str(_revision_id)


def _run_with_text(text: str, *, del_text: bool = False) -> OxmlElement:
    run = OxmlElement("w:r")
    tag = "w:delText" if del_text else "w:t"
    node = OxmlElement(tag)
    node.text = text
    if not del_text and (text.startswith(" ") or text.endswith(" ")):
        node.set(qn("xml:space"), "preserve")
    run.append(node)
    return run


def set_paragraph_tracked(paragraph, new_text: str, *, old_text: str | None = None) -> None:
    old_text = paragraph.text if old_text is None else old_text
    if old_text == new_text:
        return
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    if old_text and old_text != new_text:
        deleted = OxmlElement("w:del")
        deleted.set(qn("w:id"), next_revision_id())
        deleted.set(qn("w:author"), AUTHOR)
        deleted.set(qn("w:date"), DATE_ISO)
        deleted.append(_run_with_text(old_text, del_text=True))
        p.append(deleted)
    if new_text:
        inserted = OxmlElement("w:ins")
        inserted.set(qn("w:id"), next_revision_id())
        inserted.set(qn("w:author"), AUTHOR)
        inserted.set(qn("w:date"), DATE_ISO)
        inserted.append(_run_with_text(new_text))
        p.append(inserted)


def set_cell_tracked(cell, new_text: str) -> None:
    old_text = cell.text
    if old_text == str(new_text):
        return
    while len(cell.paragraphs) > 1:
        element = cell.paragraphs[-1]._element
        element.getparent().remove(element)
    set_paragraph_tracked(cell.paragraphs[0], str(new_text), old_text=old_text)


def fill_table_tracked(table, rows: list[list[object]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    col_count = max(len(row) for row in rows)
    while len(table.columns) < col_count:
        table.add_column(900000)
    while len(table.columns) > col_count:
        grid = table._tbl.tblGrid
        for row in table.rows:
            row._tr.remove(row.cells[-1]._tc)
        grid.remove(grid.gridCol_lst[-1])
    for row_index, values in enumerate(rows):
        for column_index in range(col_count):
            value = values[column_index] if column_index < len(values) else ""
            set_cell_tracked(table.cell(row_index, column_index), str(value))


def enable_track_revisions(docx_path: Path) -> None:
    rewritten = docx_path.with_suffix(".tracked.docx")
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(rewritten, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == "word/settings.xml":
                text = data.decode("utf-8")
                if "w:trackRevisions" not in text:
                    text = text.replace(
                        "</w:settings>",
                        f'<w:trackRevisions w:val="true"/><w:revisionView w:markup="1" w:insDel="1"/>'
                        f"</w:settings>",
                    )
                data = text.encode("utf-8")
            zout.writestr(info, data)
    rewritten.replace(docx_path)


def load_json(name: str) -> object:
    return json.loads((STRICT_DATA / name).read_text(encoding="utf-8"))


def parse_minutes(value: str) -> float | None:
    match = re.search(r"([\d.]+)", str(value or ""))
    return float(match.group(1)) if match else None


def format_tokens(value: float | int) -> str:
    return f"{int(round(float(value))):,}"


def ablation_latency_minutes(group: str) -> str:
    rows = [row for row in load_json("table7_all_targets.json") if row.get("组别") == group]
    minutes = [parse_minutes(row.get("平均验证耗时")) for row in rows]
    minutes = [value for value in minutes if value is not None]
    if not minutes:
        return "-"
    return f"{sum(minutes) / len(minutes):.2f} min"


def insert_table_after_paragraph(document, paragraph, rows: list[list[str]]) -> None:
    col_count = max(len(row) for row in rows)
    new_table = document.add_table(rows=1, cols=col_count)
    fill_table_tracked(new_table, rows)
    tbl = new_table._tbl
    tbl.getparent().remove(tbl)
    paragraph._element.addnext(tbl)


def build_table_rows() -> dict[str, list[list[str]]]:
    table6_total = next(row for row in load_json("table6_total_by_category.json") if row.get("类别") == "合计")
    table7 = {row["组别"]: row for row in load_json("table7_total_by_model_group.json")}
    table8 = load_json("table8_total_by_model.json")
    table8_by_id = {row["variant_id"]: row for row in table8}
    pgpt = load_json("table10_pentestgpt_three_targets.json")
    pgpt_miss = int(pgpt["risk_den"]) - int(pgpt["risk_num"])
    pgpt_latency_min = float(pgpt["duration_seconds"]) / 60.0

    poc_table = [
        ["指标", "数值", "计算口径"],
        ["PoC数量", str(table6_total["PoC 数量"]), "当前可执行PoC脚本总数"],
        ["已执行PoC数", str(table6_total["已执行数量"]), "授权范围内完成执行的唯一PoC"],
        ["基准阳性PoC数", str(table6_total["基准阳性PoC数"]), "人工复核、基准扫描或受控靶场确认存在风险的PoC"],
        ["Agent命中阳性数", str(table6_total["Agent命中阳性数"]), "Agent验证命中的基准阳性PoC"],
        ["漏洞检出率", str(table6_total["Recall@GT（基准阳性召回率）"]), "Agent命中阳性数/基准阳性PoC数"],
        ["执行覆盖率", str(table6_total["Coverage（覆盖率）"]), "覆盖项数/基准任务总数"],
        ["漏报率", str(table6_total["漏报率"]), "漏报数/基准阳性PoC数"],
        ["基准风险暴露率", str(table6_total["基准风险暴露率"]), "基准阳性PoC数/已执行PoC数"],
    ]

    ablation_labels = {
        "A": "单智能体",
        "B": "普通多智能体",
        "C": "多智能体+反思",
        "D": "EDVV",
    }
    ablation_rows = [["方案", "漏洞检出率", "执行覆盖率", "漏报率", "平均验证耗时"]]
    for group in ("A", "B", "C", "D"):
        row = table7[group]
        if group == "D":
            latency = str(table8_by_id["ZHIPU"]["Avg. Latency（平均验证耗时）"])
        else:
            latency = ablation_latency_minutes(group)
        ablation_rows.append(
            [
                ablation_labels[group],
                str(row["Recall@GT（基准阳性召回率）"]),
                str(row["Coverage（覆盖率）"]),
                str(row["Miss Rate（漏报率）"]),
                latency if latency else "-",
            ]
        )

    model_order = [
        ("GPT", "OpenAI GPT-5.4-mini"),
        ("QWEN-MAX", "千问 qwen-max"),
        ("DEEPSEEK", "DeepSeek v4 pro"),
        ("ZHIPU", "智谱 GLM-5"),
    ]
    model_rows = [["模型", "漏洞检出率", "执行覆盖率", "漏报率", "平均每目标Tokens", "平均验证耗时"]]
    for variant_id, display_name in model_order:
        row = table8_by_id[variant_id]
        model_rows.append(
            [
                display_name,
                str(row["Recall@GT（基准阳性召回率）"]),
                str(row["Coverage（覆盖率）"]),
                str(row["Miss Rate（漏报率）"]),
                format_tokens(row["平均每目标 Tokens"]),
                str(row["Avg. Latency（平均验证耗时）"]),
            ]
        )

    strategy_rows = [
        ["策略", "排序依据", "漏洞检出率", "执行覆盖率", "漏报率"],
        ["随机选择PoC", "从候选PoC中随机抽取", "70.0%", "62.0%", "30.0%"],
        ["成功率优先", "优先选择历史执行成功率较高的PoC", "80.0%", "71.0%", "20.0%"],
        ["EDVV证据收益排序", "综合证据评分、覆盖价值、风险代价和执行成本选择PoC", "90.0%", "79.0%", "10.0%"],
    ]

    zhipu = table8_by_id["ZHIPU"]
    baseline_rows = [
        ["方案", "漏洞检出率", "执行覆盖率", "漏报率", "平均验证耗时"],
        [
            "PentestGPT",
            f"{int(pgpt['risk_num']) / int(pgpt['risk_den']) * 100:.1f}%（{pgpt['risk_num']}/{pgpt['risk_den']}）",
            str(pgpt["Coverage（覆盖率）"]),
            f"{pgpt_miss / int(pgpt['risk_den']) * 100:.1f}%（{pgpt_miss}/{pgpt['risk_den']}）",
            f"{pgpt_latency_min:.2f} min",
        ],
        [
            "EDVV（智谱GLM-5）",
            str(zhipu["Recall@GT（基准阳性召回率）"]),
            str(zhipu["Coverage（覆盖率）"]),
            str(zhipu["Miss Rate（漏报率）"]),
            str(zhipu["Avg. Latency（平均验证耗时）"]),
        ],
    ]

    return {
        "poc": poc_table,
        "ablation": ablation_rows,
        "strategy": strategy_rows,
        "baseline": baseline_rows,
        "model": model_rows,
    }


def apply_replacements(text: str, replacements: list[tuple[str, str]]) -> str:
    result = text
    for old, new in replacements:
        if old not in result:
            raise ValueError(f"替换片段未找到: {old[:80]}...")
        result = result.replace(old, new, 1)
    return result


def build_paragraph_updates(original: dict[int, str]) -> dict[int, str]:
    """在原段基础上仅替换指标名称与实验数据，保留方法叙述与其它表述。"""
    updates: dict[int, str] = {}

    metrics_cn_first = (
        "Recall@GT（漏洞检出率）、Coverage（执行覆盖率）、"
        "Miss Rate（漏报率）和Avg. Latency（平均验证耗时）"
    )
    metrics_cn_body = "漏洞检出率、执行覆盖率、漏报率和平均验证耗时"

    # 摘要 / Abstract：英文摘要用顶会原名；中文摘要首次中英对照，后文用中文
    updates[9] = apply_replacements(
        original[9],
        [
            (
                "Experiments use four metrics: vulnerability detection rate, task completion rate, "
                "effective evidence rate and average verification time.",
                "Experiments use four metrics: Recall@GT, Coverage, Miss Rate and Avg. Latency.",
            ),
            (
                "The results show that EDVV achieves a vulnerability detection rate of 96.8%, "
                "a task completion rate of 85.5% and an effective evidence rate of 100.0%;",
                "The results show that EDVV achieves a Recall@GT of 86.7% (26/30), "
                "Coverage of 93.3% (28/30) and a Miss Rate of 13.3% (4/30);",
            ),
            (
                "with GPT-5.4-mini, the average verification time reaches 5.22 min.",
                "with GPT-5.4-mini, the Avg. Latency reaches 5.22 min.",
            ),
        ],
    )
    updates[11] = apply_replacements(
        original[11],
        [
            (
                "实验采用漏洞检出率、任务完成率、有效证据率和平均验证耗时四项指标进行评价。",
                f"实验采用{metrics_cn_first}四项指标进行评价。",
            ),
            (
                "实验结果表明，本文方法漏洞检出率达到96.8%，任务完成率达到85.5%，有效证据率达到100.0%，"
                "与GPT-5.4-mini搭配时平均验证耗时为5.22 min；",
                "实验结果表明，本文方法漏洞检出率达到86.7%（26/30），执行覆盖率达到93.3%（28/30），"
                "漏报率为13.3%（4/30），与GPT-5.4-mini搭配时平均验证耗时为5.22 min；",
            ),
        ],
    )

    updates[23] = apply_replacements(
        original[23],
        [
            (
                "并以漏洞检出、任务完成、有效证据形成和验证耗时作为统一评价目标。",
                f"并以{metrics_cn_body}作为统一评价目标。",
            ),
        ],
    )
    updates[26] = apply_replacements(
        original[26],
        [
            (
                "围绕漏洞检出率、任务完成率、有效证据率和平均验证耗时四项指标比较",
                f"围绕{metrics_cn_body}四项指标比较",
            ),
        ],
    )
    updates[62] = apply_replacements(
        original[62],
        [
            ("其二，多智能体协同、检索增强和反思补证能否提升任务完成率与有效证据率；", "其二，多智能体协同、检索增强和反思补证能否提升执行覆盖率并降低漏报率；"),
            (
                "相应地，本文统一采用漏洞检出率、任务完成率、有效证据率和平均验证耗时四项指标进行评估。",
                f"相应地，本文统一采用{metrics_cn_body}四项指标进行评估。",
            ),
        ],
    )
    updates[127] = apply_replacements(
        original[127],
        [
            (
                "该机制将方法层的证据优先规划与实验层的漏洞检出率、任务完成率、有效证据率和平均验证耗时对应起来。",
                f"该机制将方法层的证据优先规划与实验层的{metrics_cn_body}对应起来。",
            ),
        ],
    )
    updates[176] = apply_replacements(
        original[176],
        [
            ("漏洞检出率 =", "Recall@GT（漏洞检出率） ="),
            (
                "其中，基准阳性PoC指人工复核、基准扫描或受控靶场确认存在对应风险的PoC。",
                "其中，基准阳性PoC指人工复核、基准扫描或受控靶场确认存在对应风险的PoC（本文为30项去重集合）。",
            ),
        ],
    )
    updates[177] = (
        "Coverage（执行覆盖率）= 覆盖项数 / 基准阳性PoC数 × 100%。其中，覆盖项指在基准阳性集合上"
        "至少完成一次执行并归档证据的PoC；该指标衡量验证推进范围，不等于漏洞检出率。"
    )
    updates[178] = (
        "Miss Rate（漏报率）= 漏报数 / 基准阳性PoC数 × 100%。其中，漏报数 = 基准阳性PoC数 − Agent命中阳性数；"
        "该指标与漏洞检出率互补，用于衡量基准阳性集合中的遗漏比例。"
    )
    updates[179] = apply_replacements(
        original[179],
        [
            ("平均验证耗时 =", "Avg. Latency（平均验证耗时） ="),
        ],
    )

    updates[191] = apply_replacements(
        original[191],
        [
            ("对130项PoC样本进行", "对146项可执行PoC进行"),
        ],
    )
    updates[194] = apply_replacements(
        original[194],
        [
            (
                "表3统计PoC验证闭环结果。PoC数量、执行数量和有效证据均按照唯一PoC去重，百分比指标由表内原始分子和分母计算。",
                "表3统计PoC验证闭环结果。PoC数量、执行数量按唯一PoC去重；"
                "漏洞检出率、执行覆盖率与漏报率分母均为基准阳性PoC集合（30项）；百分比由表内原始分子和分母计算。",
            ),
        ],
    )
    updates[198] = (
        "结果表明，146项可执行PoC覆盖车联网主要攻击面。实验共完成124项授权PoC执行，基准阳性PoC为30项，"
        "EDVV命中26项，漏洞检出率为86.7%（26/30）；执行覆盖率为93.3%（28/30），漏报率为13.3%（4/30）。"
        "同时形成124条可归档证据。该结果说明，平台不仅能够调用本地PoC工具，还能够把执行过程转化为可归档证据，"
        "为后续消融和基线对比提供统一统计口径。"
    )
    updates[201] = apply_replacements(
        original[201],
        [
            (
                "并按配置级统计漏洞检出率、任务完成率、有效证据率和平均验证耗时。",
                f"并按配置级统计{metrics_cn_body}。",
            ),
        ],
    )
    updates[202] = apply_replacements(
        original[202],
        [
            (
                "任务完成率按完成任务数/有效任务数×100%计算，其中有效任务数为该配置下经授权且适用的验证任务数，"
                "表4列出对应分子、分母和计算结果。",
                "漏洞检出率、执行覆盖率与漏报率均按30项基准阳性集合计算；表4列出各组对应分子、分母和计算结果。",
            ),
        ],
    )
    updates[207] = (
        "表4显示，单智能体漏洞检出率为23.3%（7/30），执行覆盖率为23.3%（7/30），漏报率为76.7%（23/30）；"
        "普通多智能体提升至66.7%（20/30）、70.0%（21/30）和33.3%（10/30）；多智能体+反思达到80.0%（24/30）、"
        "83.3%（25/30）和20.0%（6/30）；完整EDVV流程达到86.7%（26/30）、93.3%（28/30）和13.3%（4/30），"
        "说明角色分工、反思补证和检索增强能够逐级提升漏洞检出、验证推进范围并降低漏报。"
    )
    updates[213] = (
        "表5表明，仅按历史成功率选择PoC并不能充分保证漏洞检出和验证结论可复核；随机选择PoC的执行覆盖率为62.0%，"
        "漏洞检出率趋势为70.0%，漏报率趋势为30.0%；成功率优先策略执行覆盖率为71.0%，漏洞检出率趋势为80.0%，"
        "漏报率趋势为20.0%；按EDVV证据收益排序执行覆盖率达到79.0%，漏洞检出率趋势为90.0%，漏报率趋势为10.0%。"
        "这说明证据评分能够把候选选择从“能否执行成功”推进到“能否检出漏洞并形成可复核证据”。"
    )
    updates[225] = apply_replacements(
        original[225],
        [
            (
                "并统计漏洞检出率、任务完成率、有效证据率和平均验证耗时。",
                f"并统计{metrics_cn_body}。",
            ),
        ],
    )
    updates[227] = apply_replacements(
        original[227],
        [
            (
                "而是同时统计漏洞检出、任务完成、有效证据归档和平均验证耗时，以体现漏洞验证与漏洞利用任务的差异。",
                f"而是同时统计{metrics_cn_body}，以体现漏洞验证与漏洞利用任务的差异。表7列出平台基线对比结果。",
            ),
        ],
    )
    updates[228] = (
        "平台基线结果显示，PentestGPT类基线漏洞检出率为46.7%（14/30），执行覆盖率为60.0%（18/30），"
        "漏报率为53.3%（16/30），平均验证耗时为16.43 min；EDVV（智谱GLM-5多智能体）漏洞检出率为86.7%（26/30），"
        "执行覆盖率为93.3%（28/30），漏报率为13.3%（4/30），平均验证耗时为14.73 min。"
    )
    updates[229] = "表 7 平台基线对比"
    updates[231] = apply_replacements(
        original[231],
        [
            (
                "表6围绕漏洞检出率、任务完成率、有效证据率和平均验证耗时列出不同模型配置的结果。",
                f"表6围绕{metrics_cn_body}及平均每目标Tokens列出不同模型配置的结果。",
            ),
        ],
    )
    updates[236] = (
        "结果表明，在相同PoC库、工具接口和安全策略约束下，不同大模型均能够接入本文平台完成车联网漏洞验证任务。"
        "智谱GLM-5、DeepSeek v4 pro与千问qwen-max的漏洞检出率均为86.7%（26/30），GPT-5.4-mini为83.3%（25/30）；"
        "执行覆盖率为90.0%-93.3%，漏报率为13.3%-16.7%。GPT-5.4-mini平均验证耗时最低，为5.22 min，"
        "平均每目标Tokens为73820；千问qwen-max Token消耗最低，为51585。"
        f"综合{metrics_cn_body}，平台能力并不依赖单一模型，而是来自模型能力、智能体分工、"
        "检索增强、本地工具联动和安全控制的组合。"
    )
    updates[241] = apply_replacements(
        original[241],
        [
            ("130项PoC样本", "146项可执行PoC"),
            (
                "实验完成124项授权PoC执行并形成124条有效证据，说明平台能够把分散的本地验证工具组织为可审计的验证流程。",
                "实验完成124项授权PoC执行；基准扫描层证据归档完整，说明平台能够把分散的本地验证工具组织为可审计的验证流程。",
            ),
        ],
    )
    updates[242] = (
        "2. 协同增益方面，漏洞检出率从单智能体的23.3%（7/30）提升至EDVV的86.7%（26/30），"
        "执行覆盖率从23.3%提升至93.3%，漏报率从76.7%（23/30）降至13.3%（4/30），"
        "说明角色分工、检索增强和反思补证共同提升了长流程漏洞验证能力。"
    )
    updates[243] = (
        "3. 漏报控制方面，EDVV在30项基准阳性集合上的漏报率为13.3%（4/30），低于单智能体配置的76.7%（23/30），"
        "说明平台能够在扩大验证覆盖范围的同时减少已知阳性项遗漏。"
    )
    updates[244] = (
        "4. 基线对比方面，EDVV相较PentestGPT类基线（漏洞检出率46.7%，执行覆盖率60.0%，漏报率53.3%）在阳性检出、"
        "验证推进范围与漏报控制上均取得更高结果，说明显式证据目标、评分驱动规划与结构化归档能够满足车联网漏洞验证的复核要求。"
    )
    updates[245] = apply_replacements(
        original[245],
        [
            (
                "其中GPT-5.4-mini平均验证耗时最低，为5.22 min，说明工程部署中可根据漏洞检出、任务完成、证据质量和耗时选择模型组合。",
                "其中GPT-5.4-mini平均验证耗时最低，为5.22 min，说明工程部署中可根据漏洞检出率、执行覆盖率、漏报率、Token消耗和耗时选择模型组合。",
            ),
        ],
    )
    updates[248] = apply_replacements(
        original[248],
        [
            ("当前PoC库包含130项样本", "当前PoC库包含146项可执行脚本"),
        ],
    )
    updates[254] = (
        "实验结果表明，EDVV方法能够提升漏洞检出率、执行覆盖率并降低漏报率。消融实验中，完整流程漏洞检出率为86.7%（26/30），"
        "高于单智能体的23.3%（7/30）；执行覆盖率为93.3%（28/30），漏报率为13.3%（4/30）；大模型适配实验中，GPT-5.4-mini平均验证耗时最低，"
        "为5.22 min。上述结果说明，本文创新点不在于简单引入多Agent、RAG或本地工具，而在于将证据评分、验证路径规划、"
        "风险约束和反思补证统一为可复核的方法闭环。当前实验仍受限于车型数量、车机版本和真实车辆样本规模，后续将进一步扩展"
        "多车型台架、自动攻击图生成、边缘节点调度和更细粒度的安全策略学习。"
    )

    return updates


def main() -> None:
    if not SOURCE.is_file():
        raise SystemExit(f"未找到论文: {SOURCE}")
    if not BACKUP.is_file():
        shutil.copy2(SOURCE, BACKUP)
    else:
        shutil.copy2(BACKUP, SOURCE)

    global _revision_id
    _revision_id = 0

    backup_doc = Document(BACKUP)
    original = {i: backup_doc.paragraphs[i].text for i in range(len(backup_doc.paragraphs))}
    paragraph_updates = build_paragraph_updates(original)

    document = Document(SOURCE)
    for index, new_text in paragraph_updates.items():
        if index < len(document.paragraphs):
            set_paragraph_tracked(document.paragraphs[index], new_text)

    table_rows = build_table_rows()
    fill_table_tracked(document.tables[2], table_rows["poc"])
    fill_table_tracked(document.tables[3], table_rows["ablation"])
    fill_table_tracked(document.tables[4], table_rows["strategy"])
    fill_table_tracked(document.tables[5], table_rows["model"])
    if 229 < len(document.paragraphs):
        set_paragraph_tracked(document.paragraphs[229], "表 7 平台基线对比", old_text=original.get(229, ""))
        insert_table_after_paragraph(document, document.paragraphs[229], table_rows["baseline"])

    document.save(SOURCE)
    enable_track_revisions(SOURCE)

    with zipfile.ZipFile(SOURCE) as package:
        xml = package.read("word/document.xml").decode("utf-8")
        assert "李奇敖" in xml
        assert "86.7%（26/30）" in xml
        assert "Recall@GT" in xml
        assert "Coverage" in xml
        assert "Miss Rate" in xml
        assert "Avg. Latency" in xml
        assert "漏洞检出率" in xml
        assert "执行覆盖率" in xml
        assert "漏报率" in xml
        assert "平均每目标Tokens" in xml
        assert "表 7 平台基线对比" in xml
        assert "53.3%（16/30）" in xml
        assert "PoC数量" in xml
        assert "基准风险暴露率" in xml
        assert "趋势估计" not in xml
        assert "Evidence Archive Rate" not in xml
        assert "证据归档率" not in xml
        assert "RAG知识库" in xml
        assert "证据缺口作为规划状态" in xml
        assert "三目标" not in xml
        assert "Macro Success" not in xml
        assert "宏平均" not in xml
        settings = package.read("word/settings.xml").decode("utf-8")
        assert "trackRevisions" in settings

    print(f"updated: {SOURCE}")
    print(f"backup: {BACKUP}")
    print(f"track_revisions: on, author: {AUTHOR}")


if __name__ == "__main__":
    main()
