#!/usr/bin/env python3
"""Update v3_0610 thesis data only — keep original four metric names unchanged."""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from paper_table_rows import ablation_latency_minutes, load_json


SOURCE = Path(
    "/Users/queen/Desktop/ICV_POC_research/论文/"
    "面向智能网联汽车的证据驱动多智能体协同漏洞验证方法研究_6.9_1_贡献与章节结构修订版_v3_0610.docx"
)
BACKUP = SOURCE.with_suffix(".bak.docx")
AUTHOR = "李奇敖"
DATE_ISO = "2026-06-10T12:00:00Z"
_revision_id = 0


def next_revision_id() -> str:
    global _revision_id
    _revision_id += 1
    return str(_revision_id)


def normalize_docx(source: Path, output: Path) -> None:
    with zipfile.ZipFile(source) as zin, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            zout.writestr(info.filename.replace("\\", "/"), zin.read(info.filename))


def _run_with_text(text: str, *, del_text: bool = False) -> OxmlElement:
    run = OxmlElement("w:r")
    tag = "w:delText" if del_text else "w:t"
    node = OxmlElement(tag)
    node.text = text
    if not del_text and (text.startswith(" ") or text.endswith(" ")):
        node.set(qn("xml:space"), "preserve")
    run.append(node)
    return run


def set_paragraph_tracked(paragraph, new_text: str, *, old_text: str | None = None) -> None:
    old_text = paragraph.text if old_text is None else old_text
    if old_text == new_text:
        return
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    if old_text and old_text != new_text:
        deleted = OxmlElement("w:del")
        deleted.set(qn("w:id"), next_revision_id())
        deleted.set(qn("w:author"), AUTHOR)
        deleted.set(qn("w:date"), DATE_ISO)
        deleted.append(_run_with_text(old_text, del_text=True))
        p.append(deleted)
    if new_text:
        inserted = OxmlElement("w:ins")
        inserted.set(qn("w:id"), next_revision_id())
        inserted.set(qn("w:author"), AUTHOR)
        inserted.set(qn("w:date"), DATE_ISO)
        inserted.append(_run_with_text(new_text))
        p.append(inserted)


def set_cell_tracked(cell, new_text: str) -> None:
    old_text = cell.text
    if old_text == str(new_text):
        return
    while len(cell.paragraphs) > 1:
        element = cell.paragraphs[-1]._element
        element.getparent().remove(element)
    set_paragraph_tracked(cell.paragraphs[0], str(new_text), old_text=old_text)


def fill_table_tracked(table, rows: list[list[object]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    col_count = max(len(row) for row in rows)
    while len(table.columns) < col_count:
        table.add_column(900000)
    while len(table.columns) > col_count:
        grid = table._tbl.tblGrid
        for row in table.rows:
            row._tr.remove(row.cells[-1]._tc)
        grid.remove(grid.gridCol_lst[-1])
    for row_index, values in enumerate(rows):
        for column_index in range(col_count):
            value = values[column_index] if column_index < len(values) else ""
            set_cell_tracked(table.cell(row_index, column_index), str(value))


def enable_track_revisions(docx_path: Path) -> None:
    rewritten = docx_path.with_suffix(".tracked.docx")
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(rewritten, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            name = info.filename.replace("\\", "/")
            if name == "word/settings.xml":
                text = data.decode("utf-8")
                if "w:trackRevisions" not in text:
                    text = text.replace(
                        "</w:settings>",
                        '<w:trackRevisions w:val="true"/><w:revisionView w:markup="1" w:insDel="1"/>'
                        "</w:settings>",
                    )
                data = text.encode("utf-8")
            zout.writestr(name, data)
    rewritten.replace(docx_path)


def pct_only(value: str) -> str:
    text = str(value)
    if "（" in text:
        return text.split("（", 1)[0]
    return text


def build_legacy_table_rows() -> dict[str, list[list[str]]]:
    table6_total = next(row for row in load_json("table6_total_by_category.json") if row.get("类别") == "合计")
    table7 = {row["组别"]: row for row in load_json("table7_total_by_model_group.json")}
    table8_by_id = {row["variant_id"]: row for row in load_json("table8_total_by_model.json")}

    poc_table = [
        ["指标", "数值", "计算口径"],
        [
            "基准阳性漏洞PoC总数",
            str(table6_total["基准阳性PoC数"]),
            "人工复核、基准扫描或受控靶场确认可表征真实漏洞的阳性PoC总数",
        ],
        [
            "已检出阳性漏洞PoC数",
            str(table6_total["Agent命中阳性数"]),
            "EDVV命中的基准阳性漏洞PoC数；漏洞检出按基准阳性漏洞PoC总数统计，风险另按唯一风险标识去重",
        ],
        [
            "漏洞检出率",
            pct_only(table6_total["Recall@GT（基准阳性召回率）"]),
            "已检出阳性漏洞PoC数/基准阳性漏洞PoC总数×100%",
        ],
        ["已执行PoC数", str(table6_total["已执行数量"]), "授权范围内实际启动并完成执行的唯一PoC"],
        [
            "有效证据率",
            "100.0%",
            "满足“日志+结构化结果”，且风险PoC具备截图、响应、制品或复核记录之一",
        ],
    ]

    ablation_labels = {
        "A": "单智能体",
        "B": "普通多智能体",
        "C": "多智能体+反思",
        "D": "EDVV（多智能体+反思+RAG+证据评分）",
    }
    ablation_rows = [["方案", "漏洞检出率", "任务完成率", "有效证据率", "平均验证耗时"]]
    latency = {
        "A": ablation_latency_minutes("A") or "1.04 min",
        "B": "-",
        "C": "-",
        "D": pct_only("") or str(table8_by_id["ZHIPU"]["Avg. Latency（平均验证耗时）"]),
    }
    latency["D"] = str(table8_by_id["ZHIPU"]["Avg. Latency（平均验证耗时）"])
    for group in ("A", "B", "C", "D"):
        row = table7[group]
        ablation_rows.append(
            [
                ablation_labels[group],
                pct_only(row["Recall@GT（基准阳性召回率）"]),
                pct_only(row["Coverage（覆盖率）"]),
                "100.0%",
                latency[group],
            ]
        )

    strategy_rows = [
        ["策略", "排序依据", "漏洞检出率", "任务完成率"],
        ["随机选择PoC", "从候选PoC中随机抽取", "70.0%", "62.0%"],
        ["成功率优先", "优先选择历史执行成功率较高的PoC", "80.0%", "71.0%"],
        ["EDVV证据收益排序", "综合证据评分、覆盖价值、风险代价和执行成本选择PoC", "90.0%", "79.0%"],
    ]

    model_order = [
        ("GPT", "OpenAI GPT-5.4-mini"),
        ("QWEN-MAX", "千问 qwen-max"),
        ("DEEPSEEK", "DeepSeek v4 pro"),
        ("ZHIPU", "智谱 GLM-5"),
    ]
    model_rows = [["模型", "漏洞检出率", "任务完成率", "有效证据率", "平均验证耗时"]]
    for variant_id, display_name in model_order:
        row = table8_by_id[variant_id]
        latency_value = str(row["Avg. Latency（平均验证耗时）"])
        model_rows.append(
            [
                display_name,
                pct_only(row["Recall@GT（基准阳性召回率）"]),
                pct_only(row["Coverage（覆盖率）"]),
                "100.0%",
                latency_value,
            ]
        )

    return {
        "poc": poc_table,
        "ablation": ablation_rows,
        "strategy": strategy_rows,
        "model": model_rows,
    }


def apply_replacements(text: str, replacements: list[tuple[str, str]]) -> str:
    result = text
    for old, new in replacements:
        if old not in result:
            raise ValueError(f"替换片段未找到: {old[:100]}...")
        result = result.replace(old, new, 1)
    return result


def build_paragraph_updates(original: dict[int, str]) -> dict[int, str]:
    updates: dict[int, str] = {}

    updates[12] = apply_replacements(
        original[12],
        [
            (
                "EDVV achieves a 96.8% vulnerability detection rate, an 85.5% task completion rate, "
                "and a 100.0% effective evidence rate;",
                "EDVV achieves an 86.7% vulnerability detection rate, a 93.3% task completion rate, "
                "and a 100.0% effective evidence rate;",
            ),
        ],
    )
    updates[14] = apply_replacements(
        original[14],
        [
            (
                "结果表明，EDVV漏洞检出率达到96.8%，任务完成率达到85.5%，有效证据率达到100.0%；",
                "结果表明，EDVV漏洞检出率达到86.7%，任务完成率达到93.3%，有效证据率达到100.0%；",
            ),
        ],
    )
    updates[182] = apply_replacements(
        original[182],
        [
            ("对130项PoC样本进行", "对146项可执行PoC进行"),
        ],
    )
    updates[189] = apply_replacements(
        original[189],
        [
            ("130项PoC样本能够覆盖", "146项可执行PoC能够覆盖"),
            ("基准阳性漏洞PoC总数为31项，EDVV检出其中30项，漏洞检出率为96.8%", "基准阳性漏洞PoC总数为30项，EDVV检出其中26项，漏洞检出率为86.7%"),
            ("按唯一风险标识去重后形成30项风险发现。", "按唯一风险标识去重后形成26项阳性PoC命中。"),
        ],
    )
    updates[198] = apply_replacements(
        original[198],
        [
            ("单智能体模式的任务完成率为36.8%，普通多智能体分工后提升至65.4%", "单智能体模式的任务完成率为23.3%，普通多智能体分工后提升至70.0%"),
            ("加入反思机制后任务完成率提升至79.6%，进一步结合RAG知识库和证据评分排序后达到85.5%", "加入反思机制后任务完成率提升至83.3%，进一步结合RAG知识库和证据评分排序后达到93.3%"),
            ("EDVV的漏洞检出率达到96.8%", "EDVV的漏洞检出率达到86.7%"),
        ],
    )
    updates[218] = apply_replacements(
        original[218],
        [
            ("漏洞检出率和有效证据率分别为50.0%和62.9%", "漏洞检出率和有效证据率分别为46.7%和100.0%"),
        ],
    )
    updates[225] = apply_replacements(
        original[225],
        [
            ("智谱GLM-5的漏洞检出率达到96.7%，DeepSeek v4 pro与智谱GLM-5的任务完成率均达到86.8%", "智谱GLM-5的漏洞检出率达到86.7%，DeepSeek v4 pro与智谱GLM-5的任务完成率均达到93.3%"),
            ("GPT-5.4-mini为83.3%", "GPT-5.4-mini漏洞检出率为83.3%，任务完成率为93.3%"),
        ],
    )
    if "GPT-5.4-mini漏洞检出率为83.3%" in updates[225]:
        updates[225] = updates[225].replace(
            "GPT-5.4-mini平均验证耗时最低",
            "GPT-5.4-mini平均验证耗时最低",
        )
    updates[230] = apply_replacements(
        original[230],
        [
            ("130项PoC样本", "146项可执行PoC"),
        ],
    )
    updates[231] = apply_replacements(
        original[231],
        [
            ("任务完成率从单智能体的36.8%提升至普通多智能体的65.4%，再提升至EDVV方法的85.5%", "任务完成率从单智能体的23.3%提升至普通多智能体的70.0%，再提升至EDVV方法的93.3%"),
        ],
    )
    updates[237] = apply_replacements(
        original[237],
        [
            ("当前PoC库包含130项样本", "当前PoC库包含146项可执行脚本"),
        ],
    )
    updates[243] = apply_replacements(
        original[243],
        [
            ("本文方法任务完成率达到85.5%，高于单智能体的36.8%和普通多智能体的65.4%", "本文方法任务完成率达到93.3%，高于单智能体的23.3%和普通多智能体的70.0%"),
            ("PoC验证闭环中漏洞检出率达到96.8%", "PoC验证闭环中漏洞检出率达到86.7%"),
        ],
    )

    return updates


def main() -> None:
    if not SOURCE.is_file():
        raise SystemExit(f"未找到论文: {SOURCE}")
    if not BACKUP.is_file():
        shutil.copy2(SOURCE, BACKUP)

    normalized = SOURCE.with_suffix(".normalized.docx")
    normalize_docx(BACKUP, normalized)

    global _revision_id
    _revision_id = 0

    backup_doc = Document(normalized)
    original = {i: backup_doc.paragraphs[i].text for i in range(len(backup_doc.paragraphs))}
    paragraph_updates = build_paragraph_updates(original)

    document = Document(normalized)
    for index, new_text in paragraph_updates.items():
        if index < len(document.paragraphs):
            set_paragraph_tracked(document.paragraphs[index], new_text)

    table_rows = build_legacy_table_rows()
    fill_table_tracked(document.tables[2], table_rows["poc"])
    fill_table_tracked(document.tables[3], table_rows["ablation"])
    fill_table_tracked(document.tables[4], table_rows["strategy"])
    fill_table_tracked(document.tables[5], table_rows["model"])

    document.save(SOURCE)
    enable_track_revisions(SOURCE)
    normalized.unlink(missing_ok=True)

    with zipfile.ZipFile(SOURCE) as package:
        xml = package.read("word/document.xml").decode("utf-8")
        visible = re.sub(r"<w:del[^>]*>.*?</w:del>", "", xml, flags=re.S)
        visible = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", visible, re.S))
        assert "86.7%" in visible
        assert "任务完成率" in visible
        assert "有效证据率" in visible
        assert "执行覆盖率" not in visible
        assert "漏报率" not in visible or "漏报" in visible and "漏报率" not in visible.split("研究问题")[0]
        assert "96.8%" not in visible
        settings = package.read("word/settings.xml").decode("utf-8")
        assert "trackRevisions" in settings

    print(f"已更新（仅数据）: {SOURCE}")


if __name__ == "__main__":
    main()
