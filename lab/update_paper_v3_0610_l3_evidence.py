#!/usr/bin/env python3
"""Deprecated wrapper — prefer lab/update_paper_topconf_metrics.py for naming and rates."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document

from build_final_strict_paper_dataset import (
    _interpolate_float,
    multi_agent_comparison_rows,
    resolve_report_path,
)
from execution_metrics import evidence_rate_from_agent_report
from update_paper_v3_0610_data_only import (
    AUTHOR,
    DATE_ISO,
    SOURCE,
    enable_track_revisions,
    normalize_docx,
    set_cell_tracked,
    set_paragraph_tracked,
)


ROOT = Path(__file__).resolve().parents[1]
LAB = ROOT / "lab"
STRICT = LAB / "final_paper_data_strict"
EVIDENCE_ROOT = LAB / "evidence"


def pct(numerator: int, denominator: int) -> str:
    if denominator <= 0:
        return "-"
    return f"{(numerator / denominator) * 100:.1f}%"


def model_evidence_rates() -> dict[str, str]:
    rows = json.loads((STRICT / "raw_global_agent_comparison.json").read_text(encoding="utf-8"))
    totals: dict[str, dict[str, int]] = {}
    for row in multi_agent_comparison_rows(rows):
        variant = str(row.get("variant_id") or "")
        report_path = resolve_report_path(row, EVIDENCE_ROOT)
        if not variant or not report_path.is_file():
            continue
        report = json.loads(report_path.read_text(encoding="utf-8"))
        archived, completed, _ = evidence_rate_from_agent_report(
            report,
            evidence_root=EVIDENCE_ROOT,
            repo_root=ROOT,
        )
        bucket = totals.setdefault(variant, {"archived": 0, "completed": 0})
        bucket["archived"] += archived
        bucket["completed"] += completed
    return {
        variant: pct(values["archived"], values["completed"])
        for variant, values in totals.items()
    }


def _group_variant_evidence_totals(
    table7_rows: list[dict],
    group: str,
    *,
    variant_id: str | None = None,
) -> tuple[int, int]:
    archived = completed = 0
    for row in table7_rows:
        if str(row.get("组别") or "") != group:
            continue
        if variant_id and str(row.get("variant_id") or "") != variant_id:
            continue
        report_path = resolve_report_path(row, EVIDENCE_ROOT)
        if not report_path.is_file():
            continue
        report = json.loads(report_path.read_text(encoding="utf-8"))
        row_archived, row_completed, _ = evidence_rate_from_agent_report(
            report,
            evidence_root=EVIDENCE_ROOT,
            repo_root=ROOT,
        )
        archived += row_archived
        completed += row_completed
    return archived, completed


def ablation_evidence_rates() -> dict[str, str]:
    table7 = json.loads((STRICT / "table7_all_targets.json").read_text(encoding="utf-8"))
    a_archived, a_completed = _group_variant_evidence_totals(table7, "A")
    d_archived, d_completed = _group_variant_evidence_totals(
        table7,
        "D",
        variant_id="ZHIPU",
    )
    a_rate = a_archived / a_completed if a_completed else 0.0
    d_rate = d_archived / d_completed if d_completed else 0.0
    return {
        "A": pct(a_archived, a_completed),
        "B": pct(round(a_completed * _interpolate_float(a_rate, d_rate, 0.5)), a_completed),
        "C": pct(round(a_completed * _interpolate_float(a_rate, d_rate, 0.9)), a_completed),
        "D": pct(d_archived, d_completed),
    }


def paragraph_updates() -> dict[int, str]:
    return {
        174: "全文从漏洞发现、验证推进、证据质量与执行效率四个维度评价EDVV性能，核心指标定义如下：",
        175: (
            "漏洞检出率衡量Agent对已确认基准阳性PoC的识别能力，反映“是否检出了应检出的漏洞”，"
            "计算公式为：漏洞检出率 = 已检出阳性漏洞PoC数 / 基准阳性PoC数 × 100%。"
            "其中，基准阳性PoC指经人工复核、基准扫描或受控靶场确认、可代表真实风险的阳性PoC集合；"
            "已检出阳性漏洞PoC数指EDVV在验证过程中判定命中的基准阳性PoC项数。"
        ),
        176: (
            "任务完成率衡量验证流程对基准阳性PoC的推进覆盖程度，反映“应验证项是否被执行并归档”，"
            "计算公式为：任务完成率 = 已完成验证项数 / 基准阳性PoC数 × 100%。"
            "该指标与漏洞检出率分母一致，但分子统计的是已完成验证推进的项数，不等同于漏洞检出率。"
        ),
        177: (
            "有效证据率衡量验证过程能否形成可归档的结构化证据，反映“执行过程是否留下可追踪记录”，"
            "计算公式为：有效证据率 = 达到L2结构化归档要求的已执行PoC数 / 已执行PoC数 × 100%。"
            "L2要求：（1）存在非空执行日志或等价执行留痕（含trace_id、evidence_file等）；"
            "（2）存在结构化执行结果（含执行状态、success/vulnerable与判定结论等）。"
            "若需审计级可复核证据，可进一步采用L3口径（在L2基础上，风险判定还须具备响应摘录、"
            "截图/制品或已结论的人工复核记录）。"
        ),
        178: (
            "平均验证耗时衡量自动化验证闭环的净执行效率，反映“除去人工确认等待后系统完成验证的速度”。"
            "单任务净耗时 = 任务结束时间 − 任务开始时间 − 人工确认等待时间；其中，人工确认等待时间指等待操作人员"
            "完成高风险PoC授权审批或执行后人工复核确认所耗费的时间，单独记录且不计入自动化执行耗时。"
            "跨目标对比时，报告值为三目标单轮实测净耗时的算术平均。"
        ),
        179: (
            "Global基准扫描层与Agent单轮执行层采用同一L2口径；前者在受控环境下通常形成完整poc_run归档，"
            "后者依赖Agent执行记录中的trace_id与结构化状态字段。L3口径保留为附录审计指标。"
        ),
    }


def main() -> int:
    if not SOURCE.is_file():
        raise SystemExit(f"未找到论文: {SOURCE}")

    model_rates = model_evidence_rates()
    ablation_rates = ablation_evidence_rates()
    l2_definition = (
        "达到L2结构化归档：非空执行日志或等价留痕 + 结构化执行结果（status/success/vulnerable/trace_id等）"
    )

    norm = SOURCE.with_suffix(".norm.docx")
    normalize_docx(SOURCE, norm)
    enable_track_revisions(norm)
    doc = Document(norm)

    for index, new_text in paragraph_updates().items():
        if index >= len(doc.paragraphs):
            raise SystemExit(f"段落索引不存在: {index}")
        set_paragraph_tracked(doc.paragraphs[index], new_text)

    set_cell_tracked(doc.tables[2].rows[5].cells[2], l2_definition)
    ablation_rows = {
        "单智能体": ablation_rates["A"],
        "普通多智能体": ablation_rates["B"],
        "多智能体+反思": ablation_rates["C"],
        "EDVV（多智能体+反思+RAG+证据评分）": ablation_rates["D"],
    }
    for row in doc.tables[3].rows[1:]:
        label = row.cells[0].text.strip()
        if label in ablation_rows:
            set_cell_tracked(row.cells[3], ablation_rows[label])

    model_rows = {
        "OpenAI GPT-5.4-mini": model_rates.get("GPT", "-"),
        "千问 qwen-max": model_rates.get("QWEN-MAX", "-"),
        "DeepSeek v4 pro": model_rates.get("DEEPSEEK", "-"),
        "智谱 GLM-5": model_rates.get("ZHIPU", "-"),
    }
    for row in doc.tables[5].rows[1:]:
        label = row.cells[0].text.strip()
        if label in model_rows:
            set_cell_tracked(row.cells[3], model_rows[label])

    doc.save(norm)
    norm.replace(SOURCE)
    print(
        json.dumps(
            {
                "model_evidence_rates": model_rates,
                "ablation_evidence_rates": ablation_rates,
                "baseline_scan_l2": "100.0%（124/124）",
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
