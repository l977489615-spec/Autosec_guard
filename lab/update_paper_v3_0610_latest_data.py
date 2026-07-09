#!/usr/bin/env python3
"""Update v3_0610 thesis: experiment data and data descriptions only (tracked revisions)."""

from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from paper_table_rows import build_paper_table_rows


SOURCE = Path(
    "/Users/queen/Desktop/ICV_POC_research/论文/"
    "面向智能网联汽车的证据驱动多智能体协同漏洞验证方法研究_6.9_1_贡献与章节结构修订版_v3_0610.docx"
)
BACKUP = SOURCE.with_suffix(".bak.docx")
AUTHOR = "李奇敖"
DATE_ISO = "2026-06-10T12:00:00Z"

_revision_id = 0


def next_revision_id() -> str:
    global _revision_id
    _revision_id += 1
    return str(_revision_id)


def normalize_docx(source: Path, output: Path) -> None:
    with zipfile.ZipFile(source) as zin, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            zout.writestr(info.filename.replace("\\", "/"), zin.read(info.filename))


def _run_with_text(text: str, *, del_text: bool = False) -> OxmlElement:
    run = OxmlElement("w:r")
    tag = "w:delText" if del_text else "w:t"
    node = OxmlElement(tag)
    node.text = text
    if not del_text and (text.startswith(" ") or text.endswith(" ")):
        node.set(qn("xml:space"), "preserve")
    run.append(node)
    return run


def set_paragraph_tracked(paragraph, new_text: str, *, old_text: str | None = None) -> None:
    old_text = paragraph.text if old_text is None else old_text
    if old_text == new_text:
        return
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    if old_text and old_text != new_text:
        deleted = OxmlElement("w:del")
        deleted.set(qn("w:id"), next_revision_id())
        deleted.set(qn("w:author"), AUTHOR)
        deleted.set(qn("w:date"), DATE_ISO)
        deleted.append(_run_with_text(old_text, del_text=True))
        p.append(deleted)
    if new_text:
        inserted = OxmlElement("w:ins")
        inserted.set(qn("w:id"), next_revision_id())
        inserted.set(qn("w:author"), AUTHOR)
        inserted.set(qn("w:date"), DATE_ISO)
        inserted.append(_run_with_text(new_text))
        p.append(inserted)


def set_cell_tracked(cell, new_text: str) -> None:
    old_text = cell.text
    if old_text == str(new_text):
        return
    while len(cell.paragraphs) > 1:
        element = cell.paragraphs[-1]._element
        element.getparent().remove(element)
    set_paragraph_tracked(cell.paragraphs[0], str(new_text), old_text=old_text)


def fill_table_tracked(table, rows: list[list[object]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    col_count = max(len(row) for row in rows)
    while len(table.columns) < col_count:
        table.add_column(900000)
    while len(table.columns) > col_count:
        grid = table._tbl.tblGrid
        for row in table.rows:
            row._tr.remove(row.cells[-1]._tc)
        grid.remove(grid.gridCol_lst[-1])
    for row_index, values in enumerate(rows):
        for column_index in range(col_count):
            value = values[column_index] if column_index < len(values) else ""
            set_cell_tracked(table.cell(row_index, column_index), str(value))


def enable_track_revisions(docx_path: Path) -> None:
    rewritten = docx_path.with_suffix(".tracked.docx")
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(rewritten, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            name = info.filename.replace("\\", "/")
            if name == "word/settings.xml":
                text = data.decode("utf-8")
                if "w:trackRevisions" not in text:
                    text = text.replace(
                        "</w:settings>",
                        '<w:trackRevisions w:val="true"/><w:revisionView w:markup="1" w:insDel="1"/>'
                        "</w:settings>",
                    )
                data = text.encode("utf-8")
            zout.writestr(name, data)
    rewritten.replace(docx_path)


def apply_replacements(text: str, replacements: list[tuple[str, str]]) -> str:
    result = text
    for old, new in replacements:
        if old not in result:
            raise ValueError(f"替换片段未找到: {old[:100]}...")
        result = result.replace(old, new, 1)
    return result


def build_paragraph_updates(original: dict[int, str]) -> dict[int, str]:
    metrics_cn_first = (
        "Recall@GT（漏洞检出率）、Coverage（执行覆盖率）、"
        "Miss Rate（漏报率）和Avg. Latency（平均验证耗时）"
    )
    metrics_cn_body = "漏洞检出率、执行覆盖率、漏报率和平均验证耗时"

    updates: dict[int, str] = {}

    updates[12] = apply_replacements(
        original[12],
        [
            (
                "Experiments evaluate vulnerability detection rate, task completion rate, "
                "effective evidence rate, and average verification time.",
                "Experiments evaluate Recall@GT, Coverage, Miss Rate, and Avg. Latency.",
            ),
            (
                "EDVV achieves a 96.8% vulnerability detection rate, an 85.5% task completion rate, "
                "and a 100.0% effective evidence rate; GPT-5.4-mini obtains the lowest average verification time of 5.22 min.",
                "EDVV achieves a Recall@GT of 86.7% (26/30), Coverage of 93.3% (28/30), "
                "and a Miss Rate of 13.3% (4/30); GPT-5.4-mini obtains the lowest Avg. Latency of 5.22 min.",
            ),
        ],
    )
    updates[14] = apply_replacements(
        original[14],
        [
            (
                "实验采用漏洞检出率、任务完成率、有效证据率和平均验证耗时四项指标评价。",
                f"实验采用{metrics_cn_first}四项指标评价。",
            ),
            (
                "结果表明，EDVV漏洞检出率达到96.8%，任务完成率达到85.5%，有效证据率达到100.0%；"
                "GPT-5.4-mini配置的平均验证耗时最低，为5.22 min。",
                "结果表明，EDVV漏洞检出率达到86.7%（26/30），执行覆盖率达到93.3%（28/30），"
                "漏报率为13.3%（4/30）；GPT-5.4-mini配置的平均验证耗时最低，为5.22 min。",
            ),
        ],
    )
    updates[29] = apply_replacements(
        original[29],
        [
            (
                "并在授权仿真、台架和实车环境中围绕漏洞检出率、任务完成率、有效证据率和平均验证耗时进行评估。",
                f"并在授权仿真、台架和实车环境中围绕{metrics_cn_body}进行评估。",
            ),
        ],
    )
    updates[61] = apply_replacements(
        original[61],
        [
            (
                "使方法目标与后续的漏洞检出率、任务完成率、有效证据率和平均验证耗时四项指标保持一致。",
                f"使方法目标与后续的{metrics_cn_body}四项指标保持一致。",
            ),
        ],
    )
    updates[65] = apply_replacements(
        original[65],
        [
            (
                "其二，多智能体协同、检索增强和反思补证能否提升任务完成率与有效证据率；",
                "其二，多智能体协同、检索增强和反思补证能否提升执行覆盖率并降低漏报率；",
            ),
            (
                "相应地，本文统一采用漏洞检出率、任务完成率、有效证据率和平均验证耗时四项指标进行评估。",
                f"相应地，本文统一采用{metrics_cn_body}四项指标进行评估。",
            ),
        ],
    )
    updates[128] = apply_replacements(
        original[128],
        [
            (
                "该机制将方法层的证据优先规划与实验层的漏洞检出率、任务完成率、有效证据率和平均验证耗时对应起来。",
                f"该机制将方法层的证据优先规划与实验层的{metrics_cn_body}对应起来。",
            ),
        ],
    )
    updates[175] = apply_replacements(
        original[175],
        [
            (
                "其中，基准阳性漏洞PoC总数指经人工复核、基准扫描或受控靶场确认可表征真实漏洞的阳性PoC总数；"
                "已检出阳性漏洞PoC数指EDVV在验证过程中命中的基准阳性漏洞PoC数。",
                "其中，基准阳性漏洞PoC总数指经人工复核、基准扫描或受控靶场确认可表征真实漏洞的阳性PoC总数"
                "（本文为30项去重并集）；已检出阳性漏洞PoC数指EDVV在验证过程中命中的基准阳性漏洞PoC数。",
            ),
        ],
    )
    updates[176] = (
        "执行覆盖率 = 覆盖项数 / 基准任务总数 × 100%。其中，覆盖项指在30项基准阳性集合上"
        "至少完成一次执行并归档证据的PoC；该指标衡量验证推进范围，不等于漏洞检出率。"
    )
    updates[177] = (
        "漏报率 = 漏报数 / 基准阳性PoC数 × 100%。其中，漏报数 = 基准阳性PoC数 − Agent命中阳性数；"
        "该指标与漏洞检出率互补，用于衡量基准阳性集合中的遗漏比例。"
    )
    updates[178] = apply_replacements(
        original[178],
        [
            ("平均验证耗时 =", "Avg. Latency（平均验证耗时） ="),
        ],
    )
    updates[179] = (
        "上述四项指标分别对应“是否检出漏洞”“验证推进范围”“遗漏控制”和“验证是否高效”。"
        "漏洞检出率、执行覆盖率与漏报率分母均为30项基准阳性PoC去重并集；百分比由表内分子和分母计算。"
    )
    updates[182] = apply_replacements(
        original[182],
        [
            ("对130项PoC样本进行", "对146项可执行PoC进行"),
        ],
    )
    updates[185] = apply_replacements(
        original[185],
        [
            (
                "表3统计PoC验证闭环结果。PoC数量、执行数量和有效证据均按照唯一PoC统计；"
                "风险数量按照唯一风险标识去重；漏洞检出率则按照“已检出阳性漏洞PoC数/基准阳性漏洞PoC总数”计算，"
                "百分比指标由表内原始分子和分母得到。",
                "表3统计PoC验证闭环结果。PoC数量、执行数量按唯一PoC去重；"
                "漏洞检出率、执行覆盖率与漏报率分母均为基准阳性PoC集合（30项）；百分比由表内原始分子和分母计算。",
            ),
        ],
    )
    updates[189] = (
        "结果表明，146项可执行PoC能够覆盖侦察、网络服务、车内协议、无线接口、应用安全、系统配置、"
        "第三方组件和高级攻击等主要攻击面。实验共完成124项授权PoC执行，基准阳性PoC为30项，"
        "EDVV命中26项，漏洞检出率为86.7%（26/30）；执行覆盖率为93.3%（28/30），漏报率为13.3%（4/30）。"
        "同时形成124条可归档证据。该结果说明，平台不仅能够调用本地PoC工具，还能够把执行过程转化为可归档证据，"
        "为后续消融和基线对比提供统一统计口径。"
    )
    updates[192] = apply_replacements(
        original[192],
        [
            (
                "并按配置级统计漏洞检出率、任务完成率、有效证据率和平均验证耗时。",
                f"并按配置级统计{metrics_cn_body}。",
            ),
        ],
    )
    updates[193] = apply_replacements(
        original[193],
        [
            (
                "任务完成率按完成任务数/有效任务数×100%计算，其中有效任务数为该配置下经授权且适用的验证任务数，表4列出对应结果。",
                "漏洞检出率、执行覆盖率与漏报率均按30项基准阳性集合计算；表4列出各组对应分子、分母和计算结果。",
            ),
        ],
    )
    updates[198] = (
        "表4显示，单智能体漏洞检出率为23.3%（7/30），执行覆盖率为23.3%（7/30），漏报率为76.7%（23/30）；"
        "普通多智能体提升至66.7%（20/30）、70.0%（21/30）和33.3%（10/30）；多智能体+反思达到80.0%（24/30）、"
        "83.3%（25/30）和20.0%（6/30）；完整EDVV流程达到86.7%（26/30）、93.3%（28/30）和13.3%（4/30），"
        "说明角色分工、反思补证和检索增强能够逐级提升漏洞检出、验证推进范围并降低漏报。"
    )
    updates[204] = (
        "表5表明，仅按历史成功率选择PoC并不能充分保证漏洞检出和验证结论可复核；随机选择PoC的执行覆盖率为62.0%，"
        "漏洞检出率趋势为70.0%，漏报率趋势为30.0%；成功率优先策略执行覆盖率为71.0%，漏洞检出率趋势为80.0%，"
        "漏报率趋势为20.0%；按EDVV证据收益排序执行覆盖率达到79.0%，漏洞检出率趋势为90.0%，漏报率趋势为10.0%。"
        "这说明证据评分能够把候选选择从“能否执行成功”推进到“能否检出漏洞并形成可复核证据”。"
    )
    updates[216] = apply_replacements(
        original[216],
        [
            (
                "并统计漏洞检出率、任务完成率、有效证据率和平均验证耗时。",
                f"并统计{metrics_cn_body}。",
            ),
        ],
    )
    updates[218] = (
        "平台基线结果显示，PentestGPT类基线漏洞检出率为46.7%（14/30），执行覆盖率为60.0%（18/30），"
        "漏报率为53.3%（16/30），平均验证耗时为16.43 min；普通Multi-Agent基线漏洞检出率为66.7%（20/30），"
        "执行覆盖率为70.0%（21/30），漏报率为33.3%（10/30）；EDVV漏洞检出率为86.7%（26/30），"
        "执行覆盖率为93.3%（28/30），漏报率为13.3%（4/30），平均验证耗时为14.73 min。"
    )
    updates[220] = apply_replacements(
        original[220],
        [
            (
                "表6围绕漏洞检出率、任务完成率、有效证据率和平均验证耗时列出不同模型配置的结果。",
                f"表6围绕{metrics_cn_body}及平均每目标Tokens列出不同模型配置的结果。",
            ),
        ],
    )
    updates[225] = (
        "结果表明，在相同PoC库、工具接口和安全策略约束下，不同大模型均能够接入本文平台完成车联网漏洞验证任务。"
        "智谱GLM-5、DeepSeek v4 pro与千问qwen-max的漏洞检出率均为86.7%（26/30），GPT-5.4-mini为83.3%（25/30）；"
        "执行覆盖率为90.0%-93.3%，漏报率为13.3%-16.7%。GPT-5.4-mini平均验证耗时最低，为5.22 min，"
        "平均每目标Tokens为73820；千问qwen-max Token消耗最低，为51585。"
        f"综合{metrics_cn_body}，平台能力并不依赖单一模型，而是来自模型能力、智能体分工、"
        "检索增强、本地工具联动和安全控制的组合。"
    )
    updates[230] = apply_replacements(
        original[230],
        [
            ("130项PoC样本", "146项可执行PoC"),
            (
                "实验完成124项授权PoC执行并形成124条有效证据，",
                "实验完成124项授权PoC执行；基准扫描层证据归档完整，",
            ),
        ],
    )
    updates[231] = (
        "2. 协同增益方面，漏洞检出率从单智能体的23.3%（7/30）提升至EDVV的86.7%（26/30），"
        "执行覆盖率从23.3%提升至93.3%，漏报率从76.7%（23/30）降至13.3%（4/30），"
        "说明角色分工、检索增强和反思补证共同提升了长流程漏洞验证能力。"
    )
    updates[232] = (
        "3. 漏报控制方面，EDVV在30项基准阳性集合上的漏报率为13.3%（4/30），低于单智能体配置的76.7%（23/30），"
        "说明平台能够在扩大验证覆盖范围的同时减少已知阳性项遗漏。"
    )
    updates[233] = (
        "4. 基线对比方面，EDVV相较PentestGPT类基线（漏洞检出率46.7%，执行覆盖率60.0%，漏报率53.3%）"
        "在阳性检出、验证推进范围与漏报控制上均取得更高结果，说明显式证据目标、评分驱动规划与结构化归档"
        "能够满足车联网漏洞验证的复核要求。"
    )
    updates[237] = apply_replacements(
        original[237],
        [
            ("当前PoC库包含130项样本", "当前PoC库包含146项可执行脚本"),
        ],
    )
    updates[243] = (
        "实验结果表明，EDVV方法能够提升漏洞检出率、执行覆盖率并降低漏报率。消融实验中，完整流程漏洞检出率为86.7%（26/30），"
        "高于单智能体的23.3%（7/30）；执行覆盖率为93.3%（28/30），漏报率为13.3%（4/30）；"
        "在模型适配实验中，GPT-5.4-mini配置的平均验证耗时最低，为5.22 min。"
        "上述结果说明，本文创新点不在于简单引入多Agent、RAG或本地工具，而在于将证据评分、验证路径规划、"
        "风险约束和反思补证统一为可复核的方法闭环。当前实验仍受限于车型数量、车机版本和真实车辆样本规模，"
        "后续将进一步扩展多车型台架、自动攻击图生成、边缘节点调度和更细粒度的安全策略学习。"
    )

    return updates


def prepare_table_rows() -> dict[str, list[list[str]]]:
    rows = build_paper_table_rows()
    ablation = rows["ablation"]
    ablation[1][0] = "单智能体"
    ablation[2][0] = "普通多智能体"
    ablation[3][0] = "多智能体+反思"
    ablation[4][0] = "EDVV（多智能体+反思+RAG+证据评分）"
    return rows


def main() -> None:
    if not SOURCE.is_file():
        raise SystemExit(f"未找到论文: {SOURCE}")

    if not BACKUP.is_file():
        shutil.copy2(SOURCE, BACKUP)
    else:
        shutil.copy2(BACKUP, SOURCE)

    normalized = SOURCE.with_suffix(".normalized.docx")
    normalize_docx(SOURCE, normalized)

    global _revision_id
    _revision_id = 0

    backup_doc = Document(normalized)
    original = {i: backup_doc.paragraphs[i].text for i in range(len(backup_doc.paragraphs))}
    paragraph_updates = build_paragraph_updates(original)

    document = Document(normalized)
    for index, new_text in paragraph_updates.items():
        if index < len(document.paragraphs):
            set_paragraph_tracked(document.paragraphs[index], new_text)

    table_rows = prepare_table_rows()
    fill_table_tracked(document.tables[2], table_rows["poc"])
    fill_table_tracked(document.tables[3], table_rows["ablation"])
    fill_table_tracked(document.tables[4], table_rows["strategy"])
    fill_table_tracked(document.tables[5], table_rows["model"])

    document.save(SOURCE)
    enable_track_revisions(SOURCE)
    normalized.unlink(missing_ok=True)

    with zipfile.ZipFile(SOURCE) as package:
        xml = package.read("word/document.xml").decode("utf-8")
        checks = [
            "86.7%（26/30）",
            "Recall@GT",
            "执行覆盖率",
            "漏报率",
            "146项可执行PoC",
            "46.7%（14/30）",
            "李奇敖",
        ]
        for item in checks:
            assert item in xml, f"缺少预期内容: {item}"
        assert "96.8%" not in xml, "仍含旧数据 96.8%"
        assert "任务完成率" not in xml or "Cybench" in xml, "正文仍含任务完成率"
        settings = package.read("word/settings.xml").decode("utf-8")
        assert "trackRevisions" in settings

    print(f"已更新: {SOURCE}")
    print(f"备份: {BACKUP}")


if __name__ == "__main__":
    main()
